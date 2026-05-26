#!/usr/bin/env python3
"""
Professional DOCX Resume Generator for Business Analyst Plata position.
Generates clean, ATS-friendly resumes in Russian and English.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    """Add a thin horizontal line using a 1-row table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    # Set minimal height and bottom border
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove borders except bottom
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '2E5090')
    tcBorders.append(bottom)
    tcPr.append(tcBorders)
    # Set row height small
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '50')
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def create_resume(lang="RU"):
    """Create a professional resume DOCX."""
    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    # Styles
    styles = doc.styles

    # Normal style
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # Heading styles
    for i in range(1, 4):
        try:
            h = styles[f'Heading {i}']
            h.font.name = 'Calibri'
            h.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            h.font.bold = True
        except:
            pass

    # ============ HEADER ============
    if lang == "RU":
        name = "Курносенко Виталий Николаевич"
        title = "Business Analyst | Корпоративные финансы | Банковские операции | Финансовая отчетность и compliance | Международные процессы"
        location = "Локация: Шанхай, Китай"
        citizenship = "Гражданство: Россия"
        phone = "Телефон: +86 15601694273"
        email = "Email: kurnosenko@mail.ru"
        wechat = "WeChat: porohnya"
        skype = "Skype: kurnosenko_vitaly"
        languages = "Языки: русский — родной; английский — свободный рабочий; китайский (мандарин) — HSK 4"
    else:
        name = "Kurnosenko Vitaly Nikolaevich"
        title = "Business Analyst | Corporate Finance | Banking Operations | Financial Reporting & Compliance | International Processes"
        location = "Location: Shanghai, China"
        citizenship = "Citizenship: Russia"
        phone = "Phone: +86 15601694273"
        email = "Email: kurnosenko@mail.ru"
        wechat = "WeChat: porohnya"
        skype = "Skype: kurnosenko_vitaly"
        languages = "Languages: Russian — native; English — fluent working proficiency; Chinese (Mandarin) — HSK 4"

    # Name
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)

    # Title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)

    # Contact line (compact)
    contact = f"{location}  |  {citizenship}  |  {phone}  |  {email}  |  {wechat}  |  {skype}"
    p = doc.add_paragraph()
    run = p.add_run(contact)
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)

    p = doc.add_paragraph()
    run = p.add_run(languages)
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)

    add_horizontal_line(doc)

    # ============ SUMMARY ============
    if lang == "RU":
        summary_title = "Краткое резюме"
        summary_text = (
            "Senior-практик и финансовый руководитель с 15+ годами опыта в корпоративных финансах, "
            "казначействе, управленческой и финансовой отчетности, внутреннем контроле и сопровождении "
            "международных операций. Сильная экспертиза в plan-fact анализе, бюджетировании, cash flow forecasting, "
            "платежном планировании, банковских процессах, финансовом контроле и compliance (включая процедуры "
            "на основе юридических документов, регламентов и IFRS/CAS).\n\n"
            "Глубокое понимание финансовых и операционных процессов в международной среде позволяет эффективно "
            "выступать в роли Business Analyst в банковском/финтех секторе: выявление бизнес-требований, анализ и "
            "оптимизация финансовых процессов, маппинг workflows, поддержка внедрения систем (ERP/QAD), подготовка "
            "аналитических материалов для стейкхолдеров и обеспечение соответствия регуляторным и внутренним требованиям.\n\n"
            "Объединяю финансовую экспертизу, системное мышление, аналитику данных (Excel/SQL/Python) и "
            "кросс-функциональную координацию. Опыт взаимодействия с советом директоров, аудиторами, банками, "
            "контрагентами и HQ в требовательной международной среде. Готов вносить вклад в развитие продуктов, "
            "процессов и compliance в Banco Plata."
        )
    else:
        summary_title = "Professional Summary"
        summary_text = (
            "Senior practitioner and financial leader with 15+ years of experience in corporate finance, treasury, "
            "management and financial reporting, internal controls, and international operations support. Strong expertise "
            "in plan-fact analysis, budgeting, cash flow forecasting, payment planning, banking processes, financial control, "
            "and compliance (including procedures based on legal documents, regulations, and IFRS/CAS).\n\n"
            "Deep understanding of financial and operational processes in an international environment enables effective "
            "performance as a Business Analyst in the banking/fintech sector: eliciting business requirements, analyzing and "
            "optimizing financial processes, mapping workflows, supporting ERP/QAD system implementations, preparing "
            "analytical materials for stakeholders, and ensuring compliance with regulatory and internal requirements.\n\n"
            "I combine financial expertise, systems thinking, data analytics (Excel/SQL/Python), and cross-functional "
            "coordination. Experience interacting with boards of directors, auditors, banks, counterparties, and HQ in a "
            "demanding international environment. Ready to contribute to product development, processes, and compliance at Banco Plata."
        )

    h = doc.add_heading(summary_title, level=1)
    h.runs[0].font.size = Pt(12)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph(summary_text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10)

    add_horizontal_line(doc)

    # ============ KEY COMPETENCIES ============
    if lang == "RU":
        comp_title = "Ключевые компетенции (Business Analyst + Finance)"
        competencies = [
            ("Бизнес-анализ и процессы", "Выявление и документирование требований, маппинг финансовых и операционных процессов, оптимизация workflows, поддержка внедрения ERP/QAD, анализ gap'ов между бизнес-потребностями и системами"),
            ("Финансовая отчетность и планирование", "Бюджетирование, лимиты, plan-fact анализ, управленческая отчетность, cash flow прогнозы по сценариям, ежемесячное платежное планирование"),
            ("Казначейство и банковские операции", "Администрирование счетов, интернет-банк, проведение платежей, контроль документов, взаимодействие с банками, поддержка привлечения/погашения займов"),
            ("Внутренний контроль и compliance", "Финансовый контроль платежей и договоров, закупочные/договорные регламенты, процедуры ценообразования, соответствие юридическим документам, внутренним процедурам и HQ-требованиям"),
            ("IFRS / CAS и аудит", "Участие в консолидированной и неконсолидированной отчетности, маппинг, взаимодействие с внешними аудиторами, подготовка раскрытий"),
            ("Аналитика и данные", "Продвинутый Excel (сводные, формулы, VBA), SQL (MS SQL), Python для обработки финансовых данных, структурированный анализ больших массивов, автоматизация отчетности"),
            ("Международные операции и стейкхолдеры", "Координация с российскими/китайскими учредителями, совет директоров, банки, клиенты, контрагенты; поддержка капитализации, ТЭО и материалов для принятия решений"),
        ]
    else:
        comp_title = "Key Competencies (Business Analyst + Finance)"
        competencies = [
            ("Business Analysis & Processes", "Eliciting and documenting requirements, mapping financial and operational processes, optimizing workflows, supporting ERP/QAD implementations, gap analysis between business needs and systems"),
            ("Financial Reporting & Planning", "Budgeting, limits, plan-fact analysis, management reporting, scenario-based cash flow forecasts, monthly payment planning"),
            ("Treasury & Banking Operations", "Account administration, internet banking, payment execution, document control, bank interactions, support for loan attraction/repayment"),
            ("Internal Control & Compliance", "Financial control of payments and contracts, procurement/contractual regulations, pricing procedures, compliance with legal documents, internal procedures, and HQ requirements"),
            ("IFRS / CAS & Audit", "Participation in consolidated and non-consolidated reporting, mapping, interaction with external auditors, preparation of disclosures"),
            ("Analytics & Data", "Advanced Excel (pivot tables, formulas, VBA), SQL (MS SQL), Python for financial data processing, structured analysis of large datasets, reporting automation"),
            ("International Operations & Stakeholders", "Coordination with Russian/Chinese founders, board of directors, banks, clients, counterparties; support for capitalization, feasibility studies, and decision-making materials"),
        ]

    h = doc.add_heading(comp_title, level=1)
    h.runs[0].font.size = Pt(12)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)

    # Competencies table
    table = doc.add_table(rows=len(competencies), cols=2)
    table.style = 'Table Grid'
    table.autofit = True

    for i, (area, evidence) in enumerate(competencies):
        row = table.rows[i]
        # Area cell (bold, shaded)
        cell0 = row.cells[0]
        cell0.text = area
        for para in cell0.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell0, "E8EEF4")

        # Evidence cell
        cell1 = row.cells[1]
        cell1.text = evidence
        for para in cell1.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(12)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_horizontal_line(doc)

    # ============ EXPERIENCE ============
    if lang == "RU":
        exp_title = "Опыт работы"
        experiences = [
            {
                "role": "Заместитель руководителя департамента экономики и финансов",
                "company": "CRAIC (Шанхай, Китай)",
                "period": "Декабрь 2017 — настоящее время",
                "bullets": [
                    "Выполнял роль ключевого бизнес-аналитика финансовой функции: анализировал бизнес-процессы, выявлял требования к отчетности и системам, настраивал финансовые отчеты в QAD для повышения прозрачности и поддержки управленческих решений.",
                    "Разрабатывал и внедрял процедуры бюджетирования, лимитирования и plan-fact анализа, обеспечивая alignment между операционными нуждами и финансовыми возможностями.",
                    "Строил сценарии cash flow и платежные планы, выявлял риски ликвидности и предлагал оптимизации — напрямую релевантно задачам BA в банковской сфере по управлению продуктами и процессами.",
                    "Участвовал в квартальной/годовой отчетности по IFRS/CAS, взаимодействовал с аудиторами, маппировал требования к данным и процессам.",
                    "Обеспечивал финансовый контроль и compliance платежных документов, договоров и операций в соответствии с юридическими документами и внутренними регламентами.",
                    "Координировал с международными стейкхолдерами (совет директоров, акционеры, банки), готовил аналитические материалы и дорожные карты.",
                ],
            },
            {
                "role": "CFO",
                "company": "Shanghai Nine-Two-Nine Aircraft Design Limited Company (Шанхай, Китай)",
                "period": "Ноябрь 2021 — 2025",
                "bullets": [
                    "Запустил финансовую функцию с нуля для международной компании: определил и внедрил бизнес-процессы, политики закупок, согласования договоров, ценообразования, бюджетирования и финансового контроля.",
                    "Выполнял полный цикл бизнес-анализа финансовых операций: от сбора требований до настройки отчетности и процедур, обеспечивая scalability и compliance.",
                    "Разрабатывал сценарии cash flow, платежные планы и контроли, выявляя bottlenecks и предлагая улучшения процессов — опыт напрямую применим к анализу банковских продуктов и клиентских journeys.",
                    "Организовал взаимодействие с аудиторами, выбор аудиторов, договорную работу и подготовку материалов для совета директоров (ТЭО, капитализация, уставный капитал).",
                    "Поддерживал банковские операции, отношения с банками и контрагентами, обеспечивая точность и своевременность финансовых процессов.",
                ],
            },
            {
                "role": "Исполнительный директор",
                "company": "Engineering Solutions (Китай / Россия)",
                "period": "2013 — 2017",
                "bullets": [
                    "Руководил закупками, продажами и ВЭД-операциями: анализировал рыночные требования, вел переговоры по контрактам, координировал ценообразование и рентабельность в международной торговой среде.",
                    "Участвовал в финансовом планировании и операционном контроле, выявляя потребности бизнеса и транслируя их в финансовые и процессные решения.",
                ],
            },
            {
                "role": "Исполнительный директор",
                "company": "New Technologies (Москва, Россия)",
                "period": "2008 — 2013",
                "bullets": [
                    "Запускал и управлял операционной деятельностью с ответственностью за прибыльность: анализировал цепочки поставок, ценообразование, рентабельность и клиентские требования в международной среде (ЮВА).",
                    "Внедрял управленческий учет и финансовый контроль, оптимизировал процессы закупок и производства.",
                ],
            },
            {
                "role": "Более ранние роли",
                "company": "2003 — 2008",
                "period": "",
                "bullets": [
                    "Практический опыт в управленческом учете, финансовом контроле, внедрении информационных систем (финансовых, производственных, управленческих).",
                    "Работа с ВЭД, логистикой, таможенными платежами, Incoterms, калькуляцией себестоимости и международными контрактами — понимание end-to-end процессов и compliance.",
                ],
            },
        ]
    else:
        exp_title = "Work Experience"
        experiences = [
            {
                "role": "Deputy Head of Economics and Finance Department",
                "company": "CRAIC (Shanghai, China)",
                "period": "December 2017 — Present",
                "bullets": [
                    "Served as key business analyst for the finance function: analyzed business processes, elicited requirements for reporting and systems, configured financial reports in QAD to improve transparency and support management decisions.",
                    "Developed and implemented budgeting, limiting, and plan-fact analysis procedures, ensuring alignment between operational needs and financial capabilities.",
                    "Built cash flow scenarios and payment plans, identified liquidity risks, and proposed optimizations — directly relevant to BA tasks in banking for product and process management.",
                    "Participated in quarterly/annual IFRS/CAS reporting, interacted with auditors, mapped data and process requirements.",
                    "Ensured financial control and compliance of payment documents, contracts, and operations in accordance with legal documents and internal regulations.",
                    "Coordinated with international stakeholders (board of directors, shareholders, banks), prepared analytical materials and roadmaps.",
                ],
            },
            {
                "role": "CFO",
                "company": "Shanghai Nine-Two-Nine Aircraft Design Limited Company (Shanghai, China)",
                "period": "November 2021 — 2025",
                "bullets": [
                    "Launched the finance function from scratch for an international company: defined and implemented business processes, procurement policies, contract approvals, pricing, budgeting, and financial control.",
                    "Performed the full cycle of business analysis of financial operations: from requirements gathering to configuring reporting and procedures, ensuring scalability and compliance.",
                    "Developed cash flow scenarios, payment plans, and controls, identifying bottlenecks and proposing process improvements — experience directly applicable to analyzing banking products and client journeys.",
                    "Organized interaction with auditors, auditor selection, contractual work, and preparation of materials for the board of directors (feasibility studies, capitalization, charter capital).",
                    "Supported banking operations, relations with banks and counterparties, ensuring accuracy and timeliness of financial processes.",
                ],
            },
            {
                "role": "Executive Director",
                "company": "Engineering Solutions (China / Russia)",
                "period": "2013 — 2017",
                "bullets": [
                    "Led procurement, sales, and foreign trade operations: analyzed market requirements, negotiated contracts, coordinated pricing and profitability in an international trading environment.",
                    "Participated in financial planning and operational control, identifying business needs and translating them into financial and process solutions.",
                ],
            },
            {
                "role": "Executive Director",
                "company": "New Technologies (Moscow, Russia)",
                "period": "2008 — 2013",
                "bullets": [
                    "Launched and managed operations with P&L responsibility: analyzed supply chains, pricing, profitability, and client requirements in an international environment (Southeast Asia).",
                    "Implemented management accounting and financial control, optimized procurement and production processes.",
                ],
            },
            {
                "role": "Earlier Roles",
                "company": "2003 — 2008",
                "period": "",
                "bullets": [
                    "Practical experience in management accounting, financial control, implementation of information systems (financial, production, management).",
                    "Work with foreign trade, logistics, customs payments, Incoterms, cost calculation, and international contracts — understanding of end-to-end processes and compliance.",
                ],
            },
        ]

    h = doc.add_heading(exp_title, level=1)
    h.runs[0].font.size = Pt(12)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)

    for exp in experiences:
        # Role + Company
        p = doc.add_paragraph()
        run = p.add_run(f"{exp['role']} — ")
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(exp['company'])
        run2.bold = True
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)

        # Period
        if exp['period']:
            p = doc.add_paragraph()
            run = p.add_run(exp['period'])
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            p.paragraph_format.space_after = Pt(1)

        # Bullets
        for bullet in exp['bullets']:
            p = doc.add_paragraph(bullet, style='List Bullet')
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs:
                run.font.size = Pt(9.5)

    add_horizontal_line(doc)

    # ============ EDUCATION ============
    if lang == "RU":
        edu_title = "Образование"
        education = [
            "MBA, РАНХиГС, 2005-2007",
            "Кандидат физико-математических наук, Южный федеральный университет, 2000-2003",
            "Диплом по физике, Южный федеральный университет, 1994-1999",
            "Колледж радиоэлектронного приборостроения, 1991-1995",
        ]
    else:
        edu_title = "Education"
        education = [
            "MBA, RANEPA, 2005-2007",
            "PhD in Physics and Mathematics, Southern Federal University, 2000-2003",
            "Physics Diploma, Southern Federal University, 1994-1999",
            "College of Radio-Electronic Instrumentation, 1991-1995",
        ]

    h = doc.add_heading(edu_title, level=1)
    h.runs[0].font.size = Pt(12)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)

    for item in education:
        p = doc.add_paragraph(f"• {item}")
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(10)

    add_horizontal_line(doc)

    # ============ CERTIFICATIONS ============
    if lang == "RU":
        cert_title = "Сертификаты"
        certs = [
            "CAP — Certified Accountant Practitioner",
            "Курсы повышения квалификации CPA Russia: финансовый учет и отчетность по IFRS, управленческий учет, налогообложение РФ, законодательство РФ",
            "Курсы главного бухгалтера",
        ]
    else:
        cert_title = "Certifications"
        certs = [
            "CAP — Certified Accountant Practitioner",
            "CPA Russia Professional Development Courses: financial accounting and IFRS reporting, management accounting, taxation of the Russian Federation, legislation of the Russian Federation",
            "Chief Accountant Courses",
        ]

    h = doc.add_heading(cert_title, level=1)
    h.runs[0].font.size = Pt(12)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)

    for item in certs:
        p = doc.add_paragraph(f"• {item}")
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(10)

    add_horizontal_line(doc)

    # ============ FIT SECTION ============
    if lang == "RU":
        fit_title = "Соответствие роли Business Analyst в Banco Plata"
        fit_bullets = [
            "Глубокий опыт финансового и операционного анализа в регулируемой среде, понимание банковских/казначейских процессов, отчетности, compliance и взаимодействия со стейкхолдерами.",
            "Сильные навыки выявления требований, документирования процессов, gap-анализа и поддержки внедрения систем (QAD/ERP) — ключевые для BA в fintech/банке.",
            "Экспертиза в аналитике данных (Excel, SQL, Python) для подготовки insights, отчетности и автоматизации.",
            "Доказанный опыт кросс-функциональной координации и коммуникации с руководством, аудиторами, банками и международными партнерами.",
            "Готов адаптироваться к специфике продуктов Banco Plata (платежи, кредитование, treasury services) и вносить вклад в улучшение клиентского опыта, внутренних процессов и регуляторного соответствия.",
        ]
        footer = "Резюме подготовлено специально для позиции Business Analyst в Banco Plata (Мексика)."
    else:
        fit_title = "Fit for Business Analyst Role at Banco Plata"
        fit_bullets = [
            "Deep experience in financial and operational analysis in a regulated environment, understanding of banking/treasury processes, reporting, compliance, and stakeholder interaction.",
            "Strong skills in requirements elicitation, process documentation, gap analysis, and system implementation support (QAD/ERP) — key for BA in fintech/banking.",
            "Expertise in data analytics (Excel, SQL, Python) for preparing insights, reporting, and automation.",
            "Proven experience in cross-functional coordination and communication with management, auditors, banks, and international partners.",
            "Ready to adapt to the specifics of Banco Plata products (payments, lending, treasury services) and contribute to improving customer experience, internal processes, and regulatory compliance.",
        ]
        footer = "Resume prepared specifically for the Business Analyst position at Banco Plata (Mexico)."

    h = doc.add_heading(fit_title, level=1)
    h.runs[0].font.size = Pt(12)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)

    for item in fit_bullets:
        p = doc.add_paragraph(f"• {item}")
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(10)

    # Footer
    p = doc.add_paragraph()
    run = p.add_run(footer)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)

    return doc

def main():
    output_dir = os.path.dirname(os.path.abspath(__file__))

    # Russian version
    print("Generating Russian resume...")
    doc_ru = create_resume(lang="RU")
    path_ru = os.path.join(output_dir, "Business_Analyst_Plata_Kurnosenko_RU.docx")
    doc_ru.save(path_ru)
    print(f"  Saved: {path_ru}")

    # English version
    print("Generating English resume...")
    doc_en = create_resume(lang="EN")
    path_en = os.path.join(output_dir, "Business_Analyst_Plata_Kurnosenko_EN.docx")
    doc_en.save(path_en)
    print(f"  Saved: {path_en}")

    print("\nDone! Two professional DOCX resumes created.")

if __name__ == "__main__":
    main()