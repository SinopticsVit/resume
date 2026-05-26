#!/usr/bin/env python3
"""Generate a warm, professional cover letter DOCX for Banco Plata."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def main():
    doc = Document()

    # Narrow margins, A4
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Normal style
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    # Header - Name
    p = doc.add_paragraph()
    run = p.add_run('Vitaly Kurnosenko')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)

    # Contact
    p = doc.add_paragraph()
    run = p.add_run('Shanghai, China  |  +86 15601694273  |  kurnosenko@mail.ru  |  WeChat: porohnya')
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)

    # Date
    p = doc.add_paragraph('May 10, 2026')
    p.paragraph_format.space_after = Pt(8)

    # Recipient
    for line in ['Hiring Team', 'Banco Plata', 'Mexico City, Mexico']:
        p = doc.add_paragraph(line)
        if line == 'Hiring Team':
            p.runs[0].bold = True
        p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Salutation
    p = doc.add_paragraph('Dear Hiring Team,')
    p.paragraph_format.space_after = Pt(8)

    # Body
    body = [
        'I am writing to apply for the Business Analyst position at Banco Plata. More than a formal application, this letter comes from a place of genuine respect for what you are building.',

        'In a market where banking often feels impersonal and full of hidden fees, Plata stands out by returning real value to people — through cashback that actually feels like "real money," high-yield savings that help customers grow their funds daily, and a mobile app that simply works. The fact that over three million users have chosen Plata, giving it a 4.8-star rating, tells me you are not just offering another financial product. You are creating an experience people trust and enjoy. That kind of customer-first thinking is rare, and it is exactly why I would be proud to contribute.',

        'My background is in launching and running financial operations from the ground up in complex, international environments. I have built finance functions from scratch, designed payment and treasury processes, implemented ERP reporting systems, and worked closely with auditors, banks, and leadership teams to keep everything compliant and transparent. What I have learned is that behind every smooth customer experience there is a web of well-designed processes, clear data flows, and thoughtful controls. I believe my combination of hands-on financial analysis, process mapping, and cross-functional coordination would allow me to help Banco Plata continue scaling its excellent product while keeping the same human touch that makes it special.',

        'What draws me most to this role is the opportunity to work on something that directly improves people\'s daily financial lives. Whether it is refining how the app surfaces insights, supporting the rollout of new features for the credit card or savings products, or helping align internal workflows so the team can move faster — I would bring both the analytical rigor and the operational empathy needed to make a real difference.',

        'I would be grateful for the chance to discuss how my experience could support Banco Plata\'s continued growth. Thank you for the work you do and for considering my application. I truly admire the product you have created and the values behind it.'
    ]

    for text in body:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(8)

    # Closing
    p = doc.add_paragraph('Warm regards,')
    p.paragraph_format.space_before = Pt(8)

    p = doc.add_paragraph()
    run = p.add_run('Vitaly Kurnosenko')
    run.bold = True
    p.paragraph_format.space_before = Pt(12)

    # Footer
    p = doc.add_paragraph()
    run = p.add_run('Enthusiastic about joining a team that makes banking feel human again.')
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)

    # Save
    folder = os.path.dirname(os.path.abspath(__file__))
    path = os.path.join(folder, 'Cover_Letter_Business_Analyst_Plata_Kurnosenko_EN.docx')
    doc.save(path)
    print(f'Saved: {path}')

if __name__ == '__main__':
    main()