#!/usr/bin/env python3
"""Generate DOCX resume for QA Engineer (Web) at Salmon."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_horizontal_line(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for border_name in ["top", "left", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "nil")
        tcBorders.append(border)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "2E5090")
    tcBorders.append(bottom)
    tcPr.append(tcBorders)
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), "50")
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


CONTENT = {
    "name": "Vitaly Kurnosenko",
    "title": (
        "QA Engineer (Web) | Manual & API Testing | Playwright E2E | "
        "Next.js / React | SQL | Fintech"
    ),
    "contact": (
        "Remote (UTC+8)  |  Telegram: @vitaly_kur  |  "
        "Email: rikkimortycrypt@gmail.com  |  Citizenship: Russia"
    ),
    "languages": (
        "Languages: English — B2+ (working); Russian — native; Chinese — HSK 4"
    ),
    "target_title": "Target Role",
    "target": [
        "QA Engineer (Web) — quality of customer-facing web products, REST/GraphQL API, "
        "Next.js/React apps, and release confidence in fintech",
        "Format: full-time, remote (GMT+5…+10) or hybrid / relocation to Manila",
    ],
    "profile_title": "Profile",
    "profile": (
        "Engineer and analyst with 3+ years building and validating web products and APIs "
        "in fintech and compliance environments. I combine manual QA discipline (test design, "
        "edge cases, clear bug reports) with hands-on automation: Playwright E2E, Vitest + "
        "Testing Library + MSW, CI/CD in GitHub Actions, and SQL checks for data correctness.\n\n"
        "Strong understanding of Next.js rendering modes (SSG, static export, client-side flows), "
        "OAuth/OIDC auth journeys, cross-browser/device behavior, and API contracts (REST, GraphQL, "
        "OpenAPI). Earlier Data Engineer background at Yofi (USA anti-fraud platform) adds depth for "
        "database validation, integration testing, and production-minded quality in regulated products."
    ),
    "fit_title": "Fit — QA Engineer (Web) @ Salmon",
    "fit_rows": [
        (
            "3+ years web application testing",
            "Sinoptics (2025–present): QA of Next.js portal and React site — manual regression, "
            "E2E, API/UI flows; Yofi (2022–2025): API/integration quality, pytest, serverless test harnesses",
        ),
        (
            "REST API & GraphQL",
            "Tested and validated REST and GraphQL integrations; Postman collections for sanity checks; "
            "Axios / contract alignment with backend",
        ),
        (
            "Manual testing & test design",
            "Boundary values, equivalence classes, negative paths, auth/token expiry, "
            "file upload edge cases, multi-step flows",
        ),
        (
            "Test cases & reporting",
            "Acceptance criteria, reproducible steps, DevTools/network logs; Jira/Confluence documentation",
        ),
        (
            "API automation",
            "Postman for integration sanity; component/API tests with MSW and Vitest; "
            "ready to extend Newman/Supertest in CI",
        ),
        (
            "Next.js (SSR / SSG / CSR)",
            "Next.js 14 multi-domain app: landing (SSG/static export) + auth portal; "
            "separate build/deploy per domain",
        ),
        (
            "Cross-browser & cross-device",
            "Manual checks on Chrome, Firefox, Safari; responsive layouts (desktop / tablet / mobile)",
        ),
        (
            "E2E automation",
            "Playwright — critical journeys including document upload; "
            "Vitest + Testing Library for components",
        ),
        (
            "SQL for data validation",
            "PostgreSQL, MS SQL, MongoDB — validate backend state, reconciliation, pipeline outputs",
        ),
        (
            "Security basics",
            "Auth guards, input validation (Zod), awareness of injection/XSS risks in forms and API params",
        ),
        (
            "DevTools & CI/CD",
            "Network, performance, console debugging; GitHub Actions, lint-staged, Husky; GitLab-style CI/CD on data platform",
        ),
        (
            "Load / performance testing",
            "DevTools performance analysis; pipeline/load awareness from data platform; ready to deepen k6",
        ),
        (
            "Cypress",
            "Primary stack: Playwright; concepts transfer directly to Cypress E2E and visual testing",
        ),
        (
            "Fintech domain",
            "Yofi — payments-adjacent anti-fraud, order validation, webhooks; "
            "15+ years finance/treasury background",
        ),
    ],
    "skills_title": "Key Skills",
    "skills": [
        ("Manual QA", "Test design, regression/smoke, exploratory testing, UAT support, bug triage"),
        ("Web E2E", "Playwright, Vitest, Testing Library, MSW"),
        ("API", "REST, GraphQL, Postman, OpenAPI/Swagger, status codes, retries, idempotency"),
        (
            "Frontend context",
            "React, Next.js 14, TypeScript, Vite, OAuth/OIDC, forms (React Hook Form, Zod)",
        ),
        ("Data & SQL", "PostgreSQL, MS SQL, MongoDB; validation queries, consistency checks"),
        ("CI/CD & quality gates", "GitHub Actions, ESLint, Prettier, Husky, lint-staged"),
        ("Collaboration", "Jira, Confluence; clear reports for dev/design/product; async remote teams"),
        ("Domain", "Fintech, compliance, document verification, e-commerce anti-fraud"),
    ],
    "projects_title": "Selected Projects (QA-relevant)",
    "projects": [
        {
            "title": "Sinoptics — multi-domain Next.js frontend (frontend-sipoptics-ru)",
            "bullets": [
                "sinoptics.ru — public marketing site (static export, SEO, responsive UI)",
                "verify.sinoptics.ru — auth portal: OAuth, dashboard, billing, invoice upload, reports",
                "QA ownership: manual regression; Playwright E2E (incl. upload flows); Vitest + MSW; GitHub Actions CI",
                "Validated SSG/static vs authenticated client flows; cross-browser checks before bucket deploy",
            ],
        },
        {
            "title": "Sinoptics AI — corporate site (frontend-sinoptics-ai)",
            "bullets": [
                "React 19 + Vite multi-page site; accessibility and layout regression before static hosting release",
            ],
        },
        {
            "title": "Yofi — merchant/admin surfaces & platform quality",
            "bullets": [
                "Contributed to Svelte/Vue merchant portal ecosystem; Playwright E2E in internal frontend repos",
                "Serverless Robot Framework test harness; pytest in Lambda and rules repos; integration-test stacks",
                "Validated Shopify webhooks, order pipelines, and GraphQL/REST services in production fintech context",
            ],
        },
    ],
    "exp_title": "Work Experience",
    "experiences": [
        {
            "role": "QA / Full Stack (Web)",
            "company": "Sinoptics (AI / compliance)",
            "period": "March 2025 — Present · Remote",
            "intro": "",
            "bullets": [
                "Ensure quality of Next.js and React web products: manual testing, Playwright E2E, component tests, API contract checks",
                "Translate product requirements into testable scenarios and acceptance criteria; edge cases for auth, uploads, billing, reports",
                "Debug issues with DevTools and backend/SQL validation; collaborate with developers on fixes and regression coverage",
                "Support release discipline: CI checks, reproducible builds, static deploy validation",
            ],
        },
        {
            "role": "Data Engineer / Integration Quality",
            "company": "Yofi (fintech / AI, USA)",
            "period": "February 2022 — March 2025 · Remote",
            "intro": "",
            "bullets": [
                "Built and maintained data/API integrations (GraphQL, REST, webhooks); Postman and automated checks for service behavior",
                "Implemented pipeline and data quality tests (dbt tests, pytest, Airflow validation); reduced incidents from bad data",
                "CI/CD and DevOps/DataOps practices; code review culture; documentation for internal API consumers (GitBook, OpenAPI-adjacent)",
                "Fintech context: order validation, fraud-adjacent flows, enterprise merchant reliability (Shopify, high-volume events)",
            ],
        },
        {
            "role": "Business / Systems Analysis",
            "company": "Sinoptics & earlier finance roles",
            "period": "2025 — present (Sinoptics BA overlap); 2003–2022 (finance/IT)",
            "intro": "",
            "bullets": [
                "Requirements → test scenarios; UAT-style validation; Postman sanity checks for integrations",
                "15+ years in corporate finance and banking operations — strong domain for regulated fintech QA",
            ],
        },
    ],
    "edu_title": "Education",
    "education": [
        "MBA, RANEPA (CIO track), 2007",
        "PhD (Candidate of Physical and Mathematical Sciences), Southern Federal University, 2003",
        "Specialist, Radiophysics, Southern Federal University, 1999",
    ],
    "extra_title": "Additional",
    "extra": [
        "Time zone: UTC+8 — aligned with Salmon core hours (Manila time)",
        "Strengths: systematic thinking, attention to edge cases, clear written communication, ownership of quality outcomes",
        "Open to relocation to Manila if required",
    ],
    "footer": (
        "Resume prepared for QA Engineer (Web) at Salmon — careers.salmon.ph"
    ),
}


def create_resume():
    c = CONTENT
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    p = doc.add_paragraph()
    run = p.add_run(c["name"])
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(c["title"])
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    run = p.add_run(c["contact"])
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(c["languages"])
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    add_horizontal_line(doc)

    h = doc.add_heading(c["target_title"], level=1)
    h.runs[0].font.size = Pt(12)
    for item in c["target"]:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10)
    add_horizontal_line(doc)

    h = doc.add_heading(c["profile_title"], level=1)
    h.runs[0].font.size = Pt(12)
    p = doc.add_paragraph(c["profile"])
    for run in p.runs:
        run.font.size = Pt(10)
    add_horizontal_line(doc)

    h = doc.add_heading(c["fit_title"], level=1)
    h.runs[0].font.size = Pt(12)
    table = doc.add_table(rows=len(c["fit_rows"]), cols=2)
    table.style = "Table Grid"
    for i, (req, evidence) in enumerate(c["fit_rows"]):
        c0 = table.rows[i].cells[0]
        c0.text = req
        for para in c0.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(c0, "E8EEF4")
        c1 = table.rows[i].cells[1]
        c1.text = evidence
        for para in c1.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
    for row in table.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(12)
    doc.add_paragraph()
    add_horizontal_line(doc)

    h = doc.add_heading(c["skills_title"], level=1)
    h.runs[0].font.size = Pt(12)
    table = doc.add_table(rows=len(c["skills"]), cols=2)
    table.style = "Table Grid"
    for i, (area, detail) in enumerate(c["skills"]):
        c0 = table.rows[i].cells[0]
        c0.text = area
        for para in c0.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(c0, "E8EEF4")
        c1 = table.rows[i].cells[1]
        c1.text = detail
        for para in c1.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
    doc.add_paragraph()
    add_horizontal_line(doc)

    h = doc.add_heading(c["projects_title"], level=1)
    h.runs[0].font.size = Pt(12)
    for project in c["projects"]:
        p = doc.add_paragraph()
        r = p.add_run(project["title"])
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)
        for bullet in project["bullets"]:
            p = doc.add_paragraph(bullet, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(9.5)
    add_horizontal_line(doc)

    h = doc.add_heading(c["exp_title"], level=1)
    h.runs[0].font.size = Pt(12)
    for exp in c["experiences"]:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{exp['role']} — ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(exp["company"])
        r2.bold = True
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)

        p = doc.add_paragraph()
        r = p.add_run(exp["period"])
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        if exp["intro"]:
            p = doc.add_paragraph(exp["intro"])
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(9.5)

        for bullet in exp["bullets"]:
            p = doc.add_paragraph(bullet, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(9.5)

    add_horizontal_line(doc)

    h = doc.add_heading(c["edu_title"], level=1)
    h.runs[0].font.size = Pt(12)
    for item in c["education"]:
        p = doc.add_paragraph(f"• {item}")
        for run in p.runs:
            run.font.size = Pt(10)

    add_horizontal_line(doc)

    h = doc.add_heading(c["extra_title"], level=1)
    h.runs[0].font.size = Pt(12)
    for item in c["extra"]:
        p = doc.add_paragraph(f"• {item}")
        for run in p.runs:
            run.font.size = Pt(10)

    add_horizontal_line(doc)

    p = doc.add_paragraph()
    r = p.add_run(c["footer"])
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def main():
    output_dir = os.path.dirname(os.path.abspath(__file__))
    filename = "QA_Engineer_Web_Salmon_Kurnosenko_EN.docx"
    path = os.path.join(output_dir, filename)
    create_resume().save(path)
    print(f"Saved: {path}")


if __name__ == "__main__":
    main()
