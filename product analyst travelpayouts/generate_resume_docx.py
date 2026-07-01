#!/usr/bin/env python3
"""Generate DOCX resume for Product Analyst (Drive) at Travelpayouts / Aviasales."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_horizontal_line(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for border_name in ["top", "left", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "nil")
        tcBorders.append(border)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "2E5090")
    tcBorders.append(bottom)
    tcPr.append(tcBorders)
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), "50")
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


CONTENT = {
    "name": "Курносенко Виталий Николаевич",
    "title": (
        "Product Analyst | SQL · Python · A/B-тесты · продуктовые метрики · "
        "дашборды · LLM"
    ),
    "contact": (
        "Локация: Азия (удалённо, любая точка мира)  |  Гражданство: Россия  |  "
        "Email: rikkimortycrypt@gmail.com  |  Telegram: @vitaly_kur"
    ),
    "languages": (
        "Языки: русский — родной; английский — B2+; китайский — HSK 4"
    ),
    "summary_title": "Профессиональное резюме",
    "summary": (
        "Продуктовый и data-аналитик с 3+ годами практики на B2C/e-commerce "
        "платформе (Yofi, США) и в AI-продукте (Sinoptics). Перевожу продуктовые "
        "и бизнес-гипотезы в метрики, дашборды, исследования и эксперименты — "
        "от формулировки задачи до рекомендаций для команды.\n\n"
        "Сильная база в математической статистике (канд. физ.-мат. наук) и "
        "инженерный контур (SQL, Python, dbt, Airflow, витрины) — умею не только "
        "считать конверсию, но и обеспечивать доверие к данным. Есть опыт "
        "A/B-тестирования ML-моделей в продакшне (shadow-predict, безопасный "
        "rollout), построения воронок и поведенческой аналитики, ROI-отчётности "
        "и работы с LLM в ежедневных задачах.\n\n"
        "Интерес к монетизации контента в мессенджерах и соцсетях (Telegram / TON). "
        "Комфортно работаю как data-партнёр продакта: самостоятельно ищу точки "
        "роста, формулирую гипотезы и довожу аналитику до понятных выводов."
    ),
    "fit_title": "Соответствие вакансии — Product Analyst (Drive), Travelpayouts",
    "fit_rows": [
        (
            "Опыт продуктового аналитика от 3 лет",
            "Yofi (2022–2025): продуктовые метрики, воронки, эксперименты, "
            "дашборды; Sinoptics (2025–н.в.): продуктовые сценарии; ранее — "
            "продуктовый аналитик (сегменты, экономика направлений)",
        ),
        (
            "A/B-тесты, полный цикл",
            "Дизайн и анализ A/B для ML-моделей (shadow-predict, маршрутизация); "
            "понимание статистики; готов углубить методы снижения дисперсии",
        ),
        (
            "Исследования и точки роста",
            "Поведенческая аналитика, кластеризация, аномалии; сегментация; "
            "ad-hoc → переиспользуемые витрины",
        ),
        (
            "Монетизация в мессенджерах и соцсетях",
            "Прямого affiliate/messenger-опыта нет; сильная база B2C-воронок, "
            "интерес и контекст Telegram / TON",
        ),
        (
            "Метрики и дашборды",
            "ROI-система, продуктовые и операционные дашборды (Power BI); "
            "готовность к Looker / Superset / Tableau",
        ),
        (
            "SQL + Python",
            "Продвинутый SQL (CTE, оконные функции); Python (pandas, PySpark)",
        ),
        (
            "LLM в работе",
            "Sinoptics: LLM-платформа, RAG, агенты; LLM для ad-hoc и валидации кода",
        ),
        (
            "Самостоятельность",
            "End-to-end от гипотезы до рекомендации; удалённая работа из Азии",
        ),
    ],
    "skills_title": "Ключевые навыки",
    "skills": [
        (
            "Продуктовая аналитика",
            "Воронки, конверсии, retention, unit-экономика; метрики под задачу; "
            "root-cause analysis",
        ),
        (
            "Эксперименты",
            "A/B-дизайн, интерпретация; shadow-тесты и controlled rollout ML-моделей",
        ),
        (
            "SQL и данные",
            "Сложные запросы, оконные функции; PostgreSQL, BigQuery, MS SQL; "
            "качество витрин",
        ),
        (
            "Python",
            "pandas, PySpark, Jupyter; автоматизация отчётности и ad-hoc",
        ),
        (
            "Data platform",
            "dbt, Airflow, Airbyte, Spark; витрины под продуктовую отчётность",
        ),
        (
            "Визуализация",
            "Power BI; понятная визуализация для продакта и стейкхолдеров",
        ),
        (
            "LLM / AI",
            "Dify, RAG, LLM-агенты; прикладное использование в analytics workflow",
        ),
        (
            "Коммуникация",
            "Jira, Confluence; объяснение сложного простым языком",
        ),
    ],
    "exp_title": "Опыт работы",
    "experiences": [
        {
            "role": "Бизнес-аналитик",
            "company": "Sinoptics (AI / compliance, удалённо)",
            "period": "март 2025 — настоящее время",
            "intro": "",
            "bullets": [
                "Продуктовые пользовательские сценарии и правила валидации: качество "
                "данных и корректность статусов в отчётности.",
                "Data flows на Python и SQL для продуктовой аналитики и AI-функций.",
                "Развёртывание LLM-платформы (Dify), RAG и мультиагентных пайплайнов; "
                "использую LLM в ежедневной аналитической работе.",
            ],
            "stack": "",
        },
        {
            "role": "Data Analyst / Analytics Engineer",
            "company": "Yofi Inc. (США, fintech / e-commerce, удалённо)",
            "period": "февраль 2022 — октябрь 2025",
            "intro": (
                "Anti-fraud и customer-intelligence платформа для мерчантов Shopify "
                "(enterprise: Lululemon)."
            ),
            "bullets": [
                "Продуктовые и операционные дашборды, ROI-система по инициативам; "
                "связь изменений с измеримым бизнес-эффектом.",
                "Анализ воронок и поведения пользователей: сессионизация, journey-метрики, "
                "кластеризация, поиск паттернов и аномалий.",
                "A/B-тестирование ML-моделей: shadow-predict, маршрутизация, "
                "безопасный rollout без риска для продакшна.",
                "Analytical datasets и витрины (dbt, BigQuery, PostgreSQL); "
                "~25 Airflow DAG; ad-hoc → переиспользуемые отчёты.",
                "Консультирование product и operations по определениям метрик.",
            ],
            "stack": (
                "Стек: SQL, Python, PySpark, dbt, Airflow, BigQuery, PostgreSQL, "
                "MongoDB, Power BI, AWS, GCP."
            ),
        },
        {
            "role": "Аналитик данных",
            "company": "ООО Формс (логистика / ВЭД)",
            "period": "январь 2017 — март 2022",
            "intro": "",
            "bullets": [
                "Структурирование больших массивов данных, сравнительный анализ, "
                "отчёты и рекомендации (SQL, Python, Power BI).",
            ],
            "stack": "",
        },
        {
            "role": "Продуктовый аналитик",
            "company": "ООО Новые технологии",
            "period": "январь 2010 — февраль 2013",
            "intro": "",
            "bullets": [
                "Анализ рынка и сегментов, конкурентной среды; бизнес-планы, "
                "запуск продуктовых линий, экономика направлений.",
            ],
            "stack": "",
        },
    ],
    "edu_title": "Образование и квалификация",
    "education": [
        "MBA, РАНХиГС, направление CIO, 2007",
        "Кандидат физико-математических наук, Южный федеральный университет, 2003",
        "Специалист, радиофизика, Южный федеральный университет, 1999",
        "CAP (Certified Accountant Practitioner); курсы CPA: IFRS, управленческий учёт",
    ],
    "extra_title": "Дополнительно",
    "extra": [
        "Опыт удалённой работы из Азии в международных командах; готов к remote-first и оплате в USD",
        "Интерес к Travelpayouts / affiliate / travel и монетизации контента в мессенджерах",
        "Сильные стороны: аналитическое мышление, самостоятельность, качество данных, английский B2+",
    ],
    "why_title": "Почему Product Analyst (Drive) в Travelpayouts",
    "why": (
        "Команда Drive строит новые направления монетизации — в том числе в "
        "мессенджерах и соцсетях. Это пересекается с моим интересом к продуктам "
        "с большой аудиторией и с опытом перевода гипотез в измеримые эксперименты "
        "и дашборды на B2C-платформе.\n\n"
        "Готов быть data-партнёром продакта: вести A/B end-to-end, искать точки "
        "роста в данных, выстраивать метрики с нуля там, где продукт только "
        "формируется, и честно объяснять ограничения данных. Формат remote-first "
        "Aviasales / Travelpayouts и работа из любой точки мира — комфортный "
        "match с моим текущим форматом."
    ),
    "footer": (
        "Резюме подготовлено для позиции Product Analyst (Drive), Travelpayouts — "
        "aviasales.ru/about/vacancies/4243702"
    ),
}


def create_resume():
    c = CONTENT
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    p = doc.add_paragraph()
    run = p.add_run(c["name"])
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(c["title"])
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    run = p.add_run(c["contact"])
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(c["languages"])
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    add_horizontal_line(doc)

    h = doc.add_heading(c["summary_title"], level=1)
    h.runs[0].font.size = Pt(12)
    p = doc.add_paragraph(c["summary"])
    for run in p.runs:
        run.font.size = Pt(10)
    add_horizontal_line(doc)

    h = doc.add_heading(c["fit_title"], level=1)
    h.runs[0].font.size = Pt(12)
    table = doc.add_table(rows=len(c["fit_rows"]), cols=2)
    table.style = "Table Grid"
    for i, (req, evidence) in enumerate(c["fit_rows"]):
        c0 = table.rows[i].cells[0]
        c0.text = req
        for para in c0.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(c0, "E8EEF4")
        c1 = table.rows[i].cells[1]
        c1.text = evidence
        for para in c1.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
    for row in table.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(12)
    doc.add_paragraph()
    add_horizontal_line(doc)

    h = doc.add_heading(c["skills_title"], level=1)
    h.runs[0].font.size = Pt(12)
    table = doc.add_table(rows=len(c["skills"]), cols=2)
    table.style = "Table Grid"
    for i, (area, detail) in enumerate(c["skills"]):
        c0 = table.rows[i].cells[0]
        c0.text = area
        for para in c0.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(c0, "E8EEF4")
        c1 = table.rows[i].cells[1]
        c1.text = detail
        for para in c1.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
    doc.add_paragraph()
    add_horizontal_line(doc)

    h = doc.add_heading(c["exp_title"], level=1)
    h.runs[0].font.size = Pt(12)
    for exp in c["experiences"]:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{exp['role']} — ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(exp["company"])
        r2.bold = True
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)

        p = doc.add_paragraph()
        r = p.add_run(exp["period"])
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        if exp["intro"]:
            p = doc.add_paragraph(exp["intro"])
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(9.5)

        for bullet in exp["bullets"]:
            p = doc.add_paragraph(bullet, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(9.5)

        if exp["stack"]:
            p = doc.add_paragraph(exp["stack"])
            for run in p.runs:
                run.font.size = Pt(9)
                run.italic = True

    add_horizontal_line(doc)

    h = doc.add_heading(c["edu_title"], level=1)
    h.runs[0].font.size = Pt(12)
    for item in c["education"]:
        p = doc.add_paragraph(f"• {item}")
        for run in p.runs:
            run.font.size = Pt(10)

    add_horizontal_line(doc)

    h = doc.add_heading(c["extra_title"], level=1)
    h.runs[0].font.size = Pt(12)
    for item in c["extra"]:
        p = doc.add_paragraph(f"• {item}")
        for run in p.runs:
            run.font.size = Pt(10)

    add_horizontal_line(doc)

    h = doc.add_heading(c["why_title"], level=1)
    h.runs[0].font.size = Pt(12)
    p = doc.add_paragraph(c["why"])
    for run in p.runs:
        run.font.size = Pt(10)

    p = doc.add_paragraph()
    r = p.add_run(c["footer"])
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def main():
    output_dir = os.path.dirname(os.path.abspath(__file__))
    filename = "Product_Analyst_Drive_Travelpayouts_Kurnosenko_RU.docx"
    path = os.path.join(output_dir, filename)
    create_resume().save(path)
    print(f"Saved: {path}")


if __name__ == "__main__":
    main()
