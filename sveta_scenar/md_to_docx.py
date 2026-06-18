"""Convert pushkin_onegin_literary_music_evening.md to Word document."""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn


def set_run_font(run, bold=False, italic=False, size=11):
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)


def add_paragraph_with_inline(doc, text, style=None, bold_prefix=False, indent=False):
    p = doc.add_paragraph(style=style)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)

    if bold_prefix and text.endswith(":"):
        parts = text.split(":", 1)
        run = p.add_run(parts[0] + ":")
        set_run_font(run, bold=True)
        if len(parts) > 1 and parts[1].strip():
            run2 = p.add_run(parts[1])
            set_run_font(run2)
        return p

    pattern = re.compile(r"\*\*(.+?)\*\*")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = p.add_run(text[pos:match.start()])
            set_run_font(run)
        run = p.add_run(match.group(1))
        set_run_font(run, bold=True)
        pos = match.end()
    if pos < len(text):
        run = p.add_run(text[pos:])
        set_run_font(run)
    return p


def add_list_item(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.5)

    pattern = re.compile(r"\*\*(.+?)\*\*")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = p.add_run(text[pos:match.start()])
            set_run_font(run)
        run = p.add_run(match.group(1))
        set_run_font(run, bold=True)
        pos = match.end()
    if pos < len(text):
        run = p.add_run(text[pos:])
        set_run_font(run)
    return p


def add_numbered_item(doc, text):
    p = doc.add_paragraph(style="List Number")
    pattern = re.compile(r"\*\*(.+?)\*\*")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = p.add_run(text[pos:match.start()])
            set_run_font(run)
        run = p.add_run(match.group(1))
        set_run_font(run, bold=True)
        pos = match.end()
    if pos < len(text):
        run = p.add_run(text[pos:])
        set_run_font(run)
    return p


def add_blockquote(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.right_indent = Cm(1.0)
        run = p.add_run(line)
        set_run_font(run, italic=True)


def convert_md_to_docx(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)

    i = 0
    quote_buffer = []

    def flush_quote():
        nonlocal quote_buffer
        if quote_buffer:
            add_blockquote(doc, quote_buffer)
            quote_buffer = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            flush_quote()
            i += 1
            continue

        if stripped.startswith("> "):
            quote_buffer.append(stripped[2:])
            i += 1
            continue

        flush_quote()

        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, bold=True, size=16)
            i += 1
            continue

        if stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=1)
            for run in p.runs:
                set_run_font(run, bold=True, size=14)
            i += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=2)
            for run in p.runs:
                set_run_font(run, bold=True, size=12)
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            add_numbered_item(doc, stripped)
            i += 1
            continue

        if stripped.startswith("- "):
            add_list_item(doc, stripped[2:])
            i += 1
            continue

        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:-2])
            set_run_font(run, bold=True, size=12)
            i += 1
            continue

        if stripped.startswith("**") and ":" in stripped:
            add_paragraph_with_inline(doc, stripped, bold_prefix=True)
            i += 1
            continue

        if stripped.endswith("  ") or (i + 1 < len(lines) and lines[i + 1].strip().startswith("Исполняет")):
            text = stripped
            while i + 1 < len(lines):
                nxt = lines[i + 1].strip()
                if not nxt:
                    break
                if nxt.startswith(("#", "-", ">", "**")) or re.match(r"^\d+\.\s", nxt):
                    break
                if nxt.startswith("Исполняет") or nxt.startswith("Фортепиано") or nxt.startswith("Романс"):
                    text += " " + nxt
                    i += 1
                    break
                text += " " + nxt
                i += 1
            add_paragraph_with_inline(doc, text)
            i += 1
            continue

        add_paragraph_with_inline(doc, stripped)
        i += 1

    flush_quote()
    doc.save(docx_path)


if __name__ == "__main__":
    base = Path(__file__).parent
    md_file = base / "pushkin_onegin_literary_music_evening.md"
    docx_file = base / "pushkin_onegin_literary_music_evening.docx"
    convert_md_to_docx(md_file, docx_file)
    print(f"Created: {docx_file}")
