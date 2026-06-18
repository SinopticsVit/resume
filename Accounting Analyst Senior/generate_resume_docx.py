#!/usr/bin/env python3
"""Generate professional DOCX resume for Sr. Accounting Analyst (Reporting) at Banco Plata."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_horizontal_line(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for border_name in ["top", "left", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "nil")
        tcBorders.append(border)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "2E5090")
    tcBorders.append(bottom)
    tcPr.append(tcBorders)
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), "50")
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


def create_resume():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    # Header
    p = doc.add_paragraph()
    run = p.add_run("Kurnosenko Vitaly Nikolaevich")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(
        "Senior Accounting Analyst (Reporting) | IFRS · Financial Close · "
        "Management Reporting | ERP · SQL · Excel"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)

    contact = (
        "Location: Shanghai, China (open to relocation to Mexico City)  |  "
        "Citizenship: Russia  |  Phone: +86 15601694273  |  "
        "Email: kurnosenko@mail.ru  |  WeChat: porohnya  |  Skype: kurnosenko_vitaly"
    )
    p = doc.add_paragraph()
    run = p.add_run(contact)
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(
        "Languages: Russian — native; English — professional working proficiency (C1 target); "
        "Chinese (Mandarin) — HSK 4"
    )
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    add_horizontal_line(doc)

    # Summary
    h = doc.add_heading("Professional Summary", level=1)
    h.runs[0].font.size = Pt(12)
    summary = (
        "Senior finance and accounting professional with 15+ years of experience in financial reporting, "
        "management reporting, plan-vs-actual and variance analysis, monthly/quarterly/annual close support, "
        "treasury and banking operations, and internal controls in international and regulated business "
        "environments. Strong hands-on background in IFRS and CAS reporting (consolidated and non-consolidated), "
        "auditor coordination, ERP-based reporting (QAD), and cash-flow and liquidity reporting.\n\n"
        "Combines deep accounting and finance expertise with SQL, Python, and reporting automation from "
        "3+ years on a production fintech/e-commerce data platform (Yofi, USA) and earlier ERP/IT leadership "
        "roles. Advanced Excel user; comfortable extracting and validating financial data, building repeatable "
        "reporting workflows, and collaborating with accounting, finance, and operations teams to improve "
        "reporting quality and efficiency.\n\n"
        "Well suited for Sr. Accounting Analyst (Reporting) at Banco Plata: preparing and analyzing financial "
        "reports under international standards, supporting close cycles, delivering management reports and "
        "ad hoc analyses, and contributing SQL-driven automation in a high-tech banking/fintech environment."
    )
    p = doc.add_paragraph(summary)
    for run in p.runs:
        run.font.size = Pt(10)
    add_horizontal_line(doc)

    # Competencies
    h = doc.add_heading("Key Competencies", level=1)
    h.runs[0].font.size = Pt(12)
    competencies = [
        (
            "Financial Reporting & Close",
            "Quarterly and annual IFRS/CAS reporting; consolidated and non-consolidated statements; "
            "close support, reporting mapping, disclosure preparation, auditor interaction",
        ),
        (
            "Management Reporting & Analysis",
            "Budgets, limits, plan-vs-actual and variance analysis, scenario-based cash-flow forecasts, "
            "monthly payment planning, board and management materials",
        ),
        (
            "IFRS & International Standards",
            "IFRS/CAS reporting cycles, audit support, chart-of-accounts mapping; CPA Russia IFRS courses; "
            "ready to apply IFRS knowledge to MX GAAP and US GAAP contexts",
        ),
        (
            "ERP & Financial Systems",
            "QAD ERP reporting configuration and generation; earlier ownership of ERP and internal reporting "
            "systems; experience comparable to SAP/Oracle/NetSuite workflows",
        ),
        (
            "SQL & Reporting Automation",
            "Advanced SQL (MS SQL, BigQuery, PostgreSQL) for data extraction, reconciliation logic, and "
            "recurring report automation; Python/pandas for financial data processing",
        ),
        (
            "Excel & Office Tools",
            "Advanced Excel (pivot tables, complex formulas, structured models); Microsoft Word for formal "
            "reporting and documentation",
        ),
        (
            "Treasury & Banking Operations",
            "Account administration, internet banking, payment execution and control, bank relations, "
            "loan drawdown/repayment support",
        ),
        (
            "Controls & Compliance",
            "Financial control of payments and contracts, procurement/contract procedures, compliance with "
            "legal documents, internal policies, and HQ requirements",
        ),
        (
            "Cross-Functional Collaboration",
            "Coordination with auditors, banks, board of directors, shareholders, operations, and "
            "international stakeholders",
        ),
    ]
    table = doc.add_table(rows=len(competencies), cols=2)
    table.style = "Table Grid"
    for i, (area, evidence) in enumerate(competencies):
        c0 = table.rows[i].cells[0]
        c0.text = area
        for para in c0.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(c0, "E8EEF4")
        c1 = table.rows[i].cells[1]
        c1.text = evidence
        for para in c1.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
    for row in table.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(12)
    doc.add_paragraph()
    add_horizontal_line(doc)

    # Experience
    h = doc.add_heading("Work Experience", level=1)
    h.runs[0].font.size = Pt(12)
    experiences = [
        {
            "role": "Deputy Head of Economics and Finance Department",
            "company": "CRAIC (Shanghai, China)",
            "period": "December 2017 — Present",
            "intro": "",
            "bullets": [
                "Supported monthly, quarterly, and annual financial reporting cycles: budgets, spending limits, "
                "plan-vs-actual analysis, and management reporting packages for leadership.",
                "Participated in quarterly and annual IFRS/CAS reporting, including unconsolidated and consolidated "
                "formats; worked with external auditors on review, mapping, and disclosure requirements.",
                "Configured and produced financial reports in QAD ERP, improving report structure, consistency, "
                "and management visibility.",
                "Built cash-flow projections under multiple scenarios and monthly payment plans; supported "
                "liquidity reporting and treasury operations.",
                "Ensured accuracy and timeliness of payment documents, banking workflows, and financial controls "
                "in line with legal documentation and internal procedures.",
                "Prepared analytical materials for the board of directors, capitalization projects, and "
                "financing-related documentation.",
            ],
            "stack": "",
        },
        {
            "role": "CFO",
            "company": "Shanghai Nine-Two-Nine Aircraft Design Limited Company (Shanghai, China)",
            "period": "November 2021 — 2025",
            "intro": "",
            "bullets": [
                "Established the finance and reporting function from scratch: policies for budgeting, contract "
                "approval, pricing, and financial control.",
                "Owned quarterly and annual unconsolidated and consolidated reporting under IFRS and CAS; "
                "managed auditor selection, contracting, and audit interaction.",
                "Delivered variance analysis, plan-vs-actual reviews, and management reporting to support "
                "operational and strategic decisions.",
                "Developed cash-flow scenarios, payment plans, and controls; maintained banking relations and "
                "payment execution discipline.",
                "Prepared board papers, feasibility studies, and capitalization roadmaps with accurate "
                "financial analysis and supporting schedules.",
            ],
            "stack": "",
        },
        {
            "role": "Data Analyst / Analytics Engineer",
            "company": "Yofi Inc. (USA, remote)",
            "period": "February 2022 — October 2025",
            "intro": (
                "Fintech/e-commerce platform (Shopify merchants including enterprise retail). "
                "Applied SQL and Python to financial and operational reporting automation."
            ),
            "bullets": [
                "Wrote advanced SQL (CTEs, window functions, complex joins) on BigQuery, PostgreSQL, and "
                "MS SQL for data extraction, reconciliation, and recurring analytical reports.",
                "Built and maintained reporting datasets and marts (dbt on BigQuery); enforced data quality "
                "tests and consistent metric definitions across finance and product stakeholders.",
                "Automated ETL/ELT reporting pipelines with Airflow (~25 production DAGs); monitored accuracy, "
                "freshness, and pipeline health.",
                "Supported ROI and operational dashboards; used Python/pandas for ad hoc financial analysis "
                "and report automation.",
                "Collaborated with cross-functional teams to translate business reporting needs into repeatable, "
                "auditable data workflows — directly relevant to SQL-driven reporting automation in banking.",
            ],
            "stack": "Stack: SQL, Python, dbt, Airflow, BigQuery, PostgreSQL, MS SQL, Power BI, Excel.",
        },
        {
            "role": "Executive Director",
            "company": "Engineering Solutions (China / Russia)",
            "period": "2013 — 2017",
            "intro": "",
            "bullets": [
                "Led procurement, sales, and foreign trade operations with responsibility for financial planning, "
                "pricing, profitability analysis, and operational control.",
                "Served as Head of IT and Finance: investment analysis, budgeting, management reporting, and "
                "process automation for KPI tracking.",
            ],
            "stack": "",
        },
        {
            "role": "Executive Director / Head of IT Department",
            "company": "New Technologies / New Engineering Solution (Moscow, Russia)",
            "period": "2003 — 2013",
            "intro": "",
            "bullets": [
                "Managed operations with P&L responsibility; implemented management accounting and financial "
                "control across international supply chains.",
                "Owned ERP and reporting systems; built structured internal reporting and data-management "
                "practices for management decision-making.",
                "Gained early experience in management accounting, financial control, and implementation of "
                "financial information systems.",
            ],
            "stack": "",
        },
    ]

    for exp in experiences:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{exp['role']} — ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(exp["company"])
        r2.bold = True
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)

        p = doc.add_paragraph()
        r = p.add_run(exp["period"])
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        if exp["intro"]:
            p = doc.add_paragraph(exp["intro"])
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(9.5)

        for bullet in exp["bullets"]:
            p = doc.add_paragraph(bullet, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(9.5)

        if exp["stack"]:
            p = doc.add_paragraph(exp["stack"])
            for run in p.runs:
                run.font.size = Pt(9)
                run.italic = True

    add_horizontal_line(doc)

    # Education
    h = doc.add_heading("Education", level=1)
    h.runs[0].font.size = Pt(12)
    for item in [
        "MBA, RANEPA, 2005–2007",
        "PhD in Physics and Mathematics, Southern Federal University, 2000–2003",
        "Physics Diploma, Southern Federal University, 1994–1999",
        "College of Radio-Electronic Instrumentation, 1991–1995",
    ]:
        p = doc.add_paragraph(f"• {item}")
        for run in p.runs:
            run.font.size = Pt(10)

    add_horizontal_line(doc)

    # Certifications
    h = doc.add_heading("Certifications", level=1)
    h.runs[0].font.size = Pt(12)
    for item in [
        "CAP — Certified Accountant Practitioner",
        "CPA Russia Professional Development Courses: financial accounting and IFRS reporting, "
        "management accounting, taxation of the Russian Federation, legislation of the Russian Federation",
        "Chief Accountant Courses",
    ]:
        p = doc.add_paragraph(f"• {item}")
        for run in p.runs:
            run.font.size = Pt(10)

    add_horizontal_line(doc)

    # Role fit
    h = doc.add_heading("Fit for Sr. Accounting Analyst (Reporting) at Banco Plata", level=1)
    h.runs[0].font.size = Pt(12)
    fit_rows = [
        (
            "IFRS and MX GAAP accounting principles",
            "Extensive IFRS/CAS reporting experience; CPA Russia IFRS training; strong foundation to adapt to Mexican GAAP in a regulated bank",
        ),
        (
            "Prepare and analyze financial reports",
            "15+ years in financial and management reporting, consolidated statements, variance and plan-vs-actual analysis",
        ),
        (
            "Monthly, quarterly, annual close support",
            "Full-cycle close participation, auditor coordination, ERP reporting, payment and treasury controls",
        ),
        (
            "Management reports & ad hoc analyses",
            "Board materials, cash-flow scenarios, budgeting, liquidity and operational analytics",
        ),
        (
            "Advanced Excel & Microsoft Word",
            "Daily use of advanced Excel; formal reporting and documentation in Word",
        ),
        (
            "Professional English (C1)",
            "Fluent working English in international environments; reporting and stakeholder communication with HQ, auditors, and banks",
        ),
        (
            "SQL for extraction & automation",
            "Production SQL on BigQuery/PostgreSQL/MS SQL; dbt, Airflow, Python for automated reporting pipelines",
        ),
        (
            "US GAAP familiarity",
            "IFRS-focused background with international reporting exposure; ready to extend to US GAAP reporting requirements",
        ),
        (
            "Bank / Fintech experience",
            "Treasury, banking operations, payment controls; 3+ years on a fintech data platform with financial KPI reporting",
        ),
        (
            "ERP / financial systems",
            "QAD ERP reporting; prior ERP ownership; transferable to NetSuite, SAP, Oracle environments",
        ),
    ]
    table = doc.add_table(rows=len(fit_rows), cols=2)
    table.style = "Table Grid"
    for i, (req, evidence) in enumerate(fit_rows):
        c0 = table.rows[i].cells[0]
        c0.text = req
        for para in c0.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(c0, "E8EEF4")
        c1 = table.rows[i].cells[1]
        c1.text = evidence
        for para in c1.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
    for row in table.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(12)

    p = doc.add_paragraph()
    r = p.add_run(
        "Resume prepared specifically for the Sr. Accounting Analyst (Reporting) position at Banco Plata (Mexico)."
    )
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)

    return doc


def main():
    output_dir = os.path.dirname(os.path.abspath(__file__))
    path = os.path.join(output_dir, "Accounting_Analyst_Senior_Plata_Kurnosenko_EN.docx")
    create_resume().save(path)
    print(f"Saved: {path}")


if __name__ == "__main__":
    main()
