#!/usr/bin/env python3
"""Generate cover letter DOCX files (EN and RU) for Banco Plata Accounting Analyst role."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def add_cover_letter(doc, lang="EN"):
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    if lang == "RU":
        name = "Курносенко Виталий Николаевич"
        contact = (
            "Шанхай, Китай  |  +86 15601694273  |  kurnosenko@mail.ru  |  "
            "WeChat: porohnya  |  Skype: kurnosenko_vitaly"
        )
        date = "18 июня 2026 г."
        recipient = [
            "Команда по подбору персонала",
            "Banco Plata",
            "Мехико, Мексика",
        ]
        salutation = "Уважаемая команда по подбору персонала,"
        body = [
            "обращаюсь с откликом на позицию Sr. Accounting Analyst (Reporting) в Banco Plata. "
            "Меня привлекает не только масштаб того, что вы построили — лицензированный банк с продуктом, "
            "которому доверяют миллионы клиентов, — но и то, что качественная отчётность и финансовая "
            "дисциплина лежат в основе устойчивого роста. В fintech, который развивается быстро, точная "
            "и своевременная reporting-функция делает инновации ответственными и управляемыми.",
            "У меня более 15 лет опыта на стыке финансовой отчётности, управленческого анализа и "
            "операционных финансов в международной и регулируемой среде. Я участвовал в ежемесячных, "
            "квартальных и годовых циклах close, готовил консолидированную и неконсолидированную "
            "отчётность по IFRS/CAS, работал с внешними аудиторами, формировал управленческую "
            "отчётность, variance-анализ и прогнозы cash flow для руководства и совета директоров. "
            "Настраивал и сопровождал отчётность в ERP (QAD), вёл казначейские и банковские процессы, "
            "обеспечивал точность платежей и соблюдение внутренних контролей в соответствии с "
            "юридическими документами и требованиями HQ.",
            "Считаю, что для этой роли в Banco Plata особенно релевантно сочетание accounting-экспертизы "
            "и навыков автоматизации отчётности. Помимо финансового бэкграунда, у меня есть три года "
            "практики на production fintech data platform (Yofi, USA), где я использовал SQL, Python, "
            "dbt и Airflow для извлечения данных, контроля качества, автоматизации регулярных отчётов "
            "и поддержки operational dashboards. Этот опыт научил меня связывать требования учёта и "
            "современные data workflows — как раз то, что нужно, когда команда ищет специалиста с SQL "
            "для reporting automation и одновременно с дисциплиной IFRS и GAAP.",
            "Комфортно работаю в условиях жёстких дедлайнов, взаимодействую с командами accounting, "
            "finance и operations и умею переводить сложные цифры в понятные управленческие выводы. "
            "Имею курсы CPA Russia по IFRS, сертификат CAP и квалификацию главного бухгалтера; готов "
            "опираться на IFRS-базу и быстро адаптироваться к MX GAAP и US GAAP в контексте "
            "регулируемого банка. Сейчас базируюсь в Шанхае и открыт к релокации в Мехико при "
            "подходящем предложении.",
            "Буду рад обсудить, как мой опыт в financial reporting, close support и SQL-автоматизации "
            "может помочь Banco Plata сохранять точность и прозрачность, на которых держится ваш продукт. "
            "Спасибо за рассмотрение моей заявки.",
        ]
        closing = "С уважением,"
        sign_name = "Виталий Курносенко"
        footer = (
            "Подготовлено для позиции Sr. Accounting Analyst (Reporting) в Banco Plata (Мексика)."
        )
    else:
        name = "Vitaly Kurnosenko"
        contact = (
            "Shanghai, China  |  +86 15601694273  |  kurnosenko@mail.ru  |  "
            "WeChat: porohnya  |  Skype: kurnosenko_vitaly"
        )
        date = "June 18, 2026"
        recipient = [
            "Hiring Team",
            "Banco Plata",
            "Mexico City, Mexico",
        ]
        salutation = "Dear Hiring Team,"
        body = [
            "I am writing to apply for the Sr. Accounting Analyst (Reporting) position at Banco Plata — "
            "and honestly, this is not just another job application for me.",
            "I have been following what your team is building for a long time, and it started much earlier "
            "than Mexico. Back in Russia, I watched Tinkoff change what people expected from a bank: an app "
            "that actually works, products explained in plain language, and a feeling that the bank is on "
            "your side, not working against you. I admired that product deeply — not as a marketer's slogan, "
            "but as something you could feel in everyday life. When I learned that the same spirit was taking "
            "root in Latin America through Plata, I felt genuine excitement. A bank that earns trust through "
            "cashback people can feel, savings that help money grow, and an experience millions of users rate "
            "highly — that is rare. It reminded me why I chose finance in the first place: numbers should "
            "serve people, not confuse them.",
            "That is also why this particular role speaks to me. Reporting is often seen as \"back office,\" "
            "but in a bank like yours it is part of the product's honesty. Customers trust Plata because "
            "things work smoothly on the surface — and behind that smoothness there must be accurate reports, "
            "disciplined close processes, and people who care that the numbers are right. That is the work I "
            "have been doing for 15+ years: financial reporting and close support, IFRS/CAS consolidated and "
            "non-consolidated statements, management reporting and variance analysis, coordination with "
            "auditors, treasury and banking operations, and ERP-based reporting in QAD. I have built finance "
            "functions from scratch, worked with boards and shareholders, and learned that good reporting is "
            "not bureaucracy — it is respect for everyone who depends on those numbers.",
            "I also bring something that fits your high-tech environment: three years on a fintech data "
            "platform (Yofi, USA), where I used SQL, Python, dbt, and Airflow to automate recurring reports, "
            "validate data quality, and support operational dashboards. I love that Plata asks for both "
            "IFRS/GAAP discipline and SQL for reporting automation — that is exactly the bridge I enjoy "
            "building between accounting and modern data workflows.",
            "I hold CPA Russia IFRS training, CAP certification, and chief accountant qualifications. I am "
            "based in Shanghai and open to relocation to Mexico City. I would be grateful for the chance to "
            "talk — not only about my experience, but about why I would be proud to help protect the "
            "transparency behind a product I have admired since the Tinkoff days.",
            "Thank you for the work you do and for considering my application.",
        ]
        closing = "Warm regards,"
        sign_name = "Vitaly Kurnosenko"
        footer = (
            "Prepared for the Sr. Accounting Analyst (Reporting) role at Banco Plata (Mexico)."
        )

    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run(contact)
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph(date)
    p.paragraph_format.space_after = Pt(8)

    for i, line in enumerate(recipient):
        p = doc.add_paragraph(line)
        if i == 0:
            p.runs[0].bold = True
        p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph(salutation)
    p.paragraph_format.space_after = Pt(8)

    for text in body:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph(closing)
    p.paragraph_format.space_before = Pt(8)

    p = doc.add_paragraph()
    run = p.add_run(sign_name)
    run.bold = True
    p.paragraph_format.space_before = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(footer)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)


def main():
    output_dir = os.path.dirname(os.path.abspath(__file__))

    doc_en = Document()
    add_cover_letter(doc_en, lang="EN")
    path_en = os.path.join(
        output_dir, "Cover_Letter_Accounting_Analyst_Senior_Plata_Kurnosenko_EN.docx"
    )
    doc_en.save(path_en)
    print(f"Saved: {path_en}")

    doc_ru = Document()
    add_cover_letter(doc_ru, lang="RU")
    path_ru = os.path.join(
        output_dir, "Cover_Letter_Accounting_Analyst_Senior_Plata_Kurnosenko_RU.docx"
    )
    doc_ru.save(path_ru)
    print(f"Saved: {path_ru}")


if __name__ == "__main__":
    main()
