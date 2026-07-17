#!/usr/bin/env python3
"""Generate DOCX for System Analyst resume (RU/EN) and cover letter (EN)."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

HERE = Path(__file__).parent
sys.path.insert(0, str(HERE.parent / "analytics_engineer"))

from generate_plata_risk_docx import (  # noqa: E402
    NAVY,
    add_rich_text,
    convert_resume_md_to_docx,
    set_run_font,
)

FILES = [
    (HERE / "System_Analyst_Kurnosenko_RU.md", HERE / "System_Analyst_Kurnosenko_RU.docx"),
    (HERE / "System_Analyst_Kurnosenko_EN.md", HERE / "System_Analyst_Kurnosenko_EN.docx"),
]
COVER_MD = HERE / "Cover_Letter_System_Analyst_Kurnosenko_EN.md"
COVER_DOCX = HERE / "Cover_Letter_System_Analyst_Kurnosenko_EN.docx"


def convert_cover_md_to_docx(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    in_body = False

    for line in lines:
        stripped = line.strip()
        if not stripped or stripped.startswith("# Cover Letter") or stripped == "---":
            continue

        if stripped.startswith("**Vitaly Kurnosenko**") and not in_body:
            p = doc.add_paragraph()
            run = p.add_run("Vitaly Kurnosenko")
            set_run_font(run, bold=True, size=16, color=NAVY)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if stripped.startswith("Shanghai") or stripped.startswith("+86"):
            p = doc.add_paragraph()
            add_rich_text(p, stripped.replace("**", ""), size=9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if re.match(r"^\*\*(January|February|March|April|May|June|July|August|September|October|November|December)\b", stripped) or stripped.startswith("**Hiring"):
            p = doc.add_paragraph()
            add_rich_text(p, stripped, size=11)
            if "Hiring" in stripped:
                for run in p.runs:
                    run.bold = True
            continue

        if stripped.startswith("Dear "):
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_after = Pt(8)
            in_body = True
            continue

        if stripped.startswith("Sincerely") or stripped.startswith("Warm regards"):
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_before = Pt(12)
            in_body = False
            continue

        if stripped == "**Vitaly Kurnosenko**":
            p = doc.add_paragraph()
            run = p.add_run("Vitaly Kurnosenko")
            set_run_font(run, bold=True, size=11)
            continue

        if in_body:
            p = doc.add_paragraph()
            add_rich_text(p, stripped, size=11)
            p.paragraph_format.space_after = Pt(8)

    doc.save(docx_path)


if __name__ == "__main__":
    for md, docx in FILES:
        convert_resume_md_to_docx(md, docx)
        print("DOCX written:", docx)
    convert_cover_md_to_docx(COVER_MD, COVER_DOCX)
    print("DOCX written:", COVER_DOCX)
