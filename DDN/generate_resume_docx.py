#!/usr/bin/env python3
"""Generate professional DOCX resume for Director, Deal Desk at DDN."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_horizontal_line(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for border_name in ["top", "left", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "nil")
        tcBorders.append(border)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "2E5090")
    tcBorders.append(bottom)
    tcPr.append(tcBorders)
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), "50")
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


def _add_table(doc, rows, col0_width=5.5, col1_width=12):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (left, right) in enumerate(rows):
        c0 = table.rows[i].cells[0]
        c0.text = left
        for para in c0.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(c0, "E8EEF4")
        c1 = table.rows[i].cells[1]
        c1.text = right
        for para in c1.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
    for row in table.rows:
        row.cells[0].width = Cm(col0_width)
        row.cells[1].width = Cm(col1_width)
    doc.add_paragraph()


def create_resume():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    p = doc.add_paragraph()
    run = p.add_run("Kurnosenko Vitaly Nikolaevich")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(
        "Director, Deal Desk | Strategic Finance · Commercial Deal Structuring · "
        "Margin Analysis · Cross-Functional Deal Execution"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)

    contact = (
        "Location: Russia (Remote — EMEA time zones)  |  Citizenship: Russia  |  "
        "Phone: +86 15601694273  |  Email: kurnosenko@mail.ru  |  "
        "Telegram: @vitaly_kur  |  WeChat: porohnya"
    )
    p = doc.add_paragraph()
    run = p.add_run(contact)
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(
        "Languages: Russian — native; English — professional working proficiency; "
        "Chinese (Mandarin) — HSK 4"
    )
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    add_horizontal_line(doc)

    h = doc.add_heading("Professional Summary", level=1)
    h.runs[0].font.size = Pt(12)
    summary = (
        "Finance and commercial leader with 15+ years of experience at the intersection of "
        "Sales support, pricing, contract governance, profitability analysis, and executive "
        "reporting in international B2B environments. Built finance functions and approval "
        "frameworks from scratch; acted as a strategic partner to commercial teams on "
        "multi-party deals, margin trade-offs, and exception-based decisions.\n\n"
        "Expert-level Excel modeling (PhD in mathematical sciences, MBA, CAP) combined with "
        "hands-on ERP/QAD reporting, IFRS/CAS reporting cycles, and recent SaaS/tech analytics "
        "experience (Yofi, Sinoptics AI). Comfortable building processes while scaling — "
        "playbooks, guardrails, and cross-functional orchestration between Sales, Finance, "
        "Legal, and Operations. Time-zone resilient operator based in Russia/Asia.\n\n"
        "Seeking to bring analytical rigor, commercial judgment, and executive communication "
        "to Director, Deal Desk at DDN — driving deal velocity while protecting long-term "
        "margin and revenue integrity."
    )
    p = doc.add_paragraph(summary)
    for run in p.runs:
        run.font.size = Pt(10)
    add_horizontal_line(doc)

    h = doc.add_heading("Role Fit — Director, Deal Desk at DDN", level=1)
    h.runs[0].font.size = Pt(12)
    _add_table(
        doc,
        [
            (
                "Deal Desk / strategic finance in tech",
                "15+ years commercial finance leadership; recent SaaS platform experience "
                "(Yofi — enterprise Shopify merchants); AI product analytics (Sinoptics)",
            ),
            (
                "Deal structuring & advisory",
                "B2B contract negotiation, pricing, payment terms, profitability management; "
                "CFO-level approval policies; ROI and merchant KPI analytics",
            ),
            (
                "Commercial guardrails",
                "Designed contract approval, pricing, budgeting, and financial control procedures; "
                "board-ready exception materials",
            ),
            (
                "Cross-functional orchestration",
                "Bridge between sales/commercial, finance, legal, auditors, banks, board in "
                "international JV environments",
            ),
            (
                "CPQ & process optimization",
                "ERP/QAD reporting setup; process formalization from greenfield; reporting "
                "automation (SQL, Python, Power BI)",
            ),
            (
                "Margin analysis & modeling",
                "Expert Excel; plan-vs-actual, scenario cash-flow models, cost/profitability "
                "analysis; PhD-level quantitative foundation",
            ),
            (
                "Negotiation & executive presence",
                "B2B negotiations with enterprise clients; board papers, capitalization "
                "roadmaps, audit and financing discussions",
            ),
            (
                "Build Deal Desk function",
                "Launched finance function from scratch (CFO); defined policies where none existed",
            ),
            (
                "EMEA / time-zone resilience",
                "Russia (Remote); 7+ years leading China-based international operations; "
                "independent cross-functional execution",
            ),
        ],
    )
    add_horizontal_line(doc)

    h = doc.add_heading("Core Skills", level=1)
    h.runs[0].font.size = Pt(12)
    _add_table(
        doc,
        [
            (
                "Deal & Commercial Finance",
                "Deal profitability, pricing and margin trade-offs, payment terms, multi-scenario "
                "modeling, exception business cases, contract governance",
            ),
            (
                "Financial Modeling",
                "Expert Excel; plan-vs-actual, cash-flow projections; SQL/Python for analysis",
            ),
            (
                "Process & Controls",
                "Approval matrices, financial control, budgeting limits, IFRS/CAS, audit, QAD",
            ),
            (
                "Cross-Functional Leadership",
                "Sales partnership, legal-document execution, board materials, stakeholder mgmt",
            ),
            (
                "Tech & SaaS Context",
                "ROI dashboards, merchant KPI logic, remote US product teams (Yofi); AI workflows",
            ),
            (
                "Communication",
                "Complex structures in plain language for field teams and executive leadership",
            ),
        ],
        col0_width=4.5,
        col1_width=13,
    )
    add_horizontal_line(doc)

    h = doc.add_heading("Work Experience", level=1)
    h.runs[0].font.size = Pt(12)
    experiences = [
        {
            "role": "CFO",
            "company": "Shanghai Nine-Two-Nine Aircraft Design Limited Company (Shanghai, China)",
            "period": "November 2021 — 2025",
            "intro": "",
            "bullets": [
                "Launched finance and commercial control from scratch: procurement policies, "
                "contract approval workflows, pricing, budgeting, and financial governance.",
                "Built Excel scenario models — budgets, limits, plan-vs-actual, cash-flow "
                "projections, and monthly payment plans.",
                "Primary finance partner on capital projects, feasibility studies, and board "
                "materials; IFRS/CAS reporting and external audit cycles.",
                "Managed banking, counterparty, and founder/board relationships; payment "
                "execution aligned with internal procedures.",
            ],
        },
        {
            "role": "Deputy Head of Economics and Finance",
            "company": "CRAIC (Shanghai, China)",
            "period": "December 2017 — July 2024",
            "intro": "",
            "bullets": [
                "Led economics and finance for a China-based international JV: budgeting, "
                "management reporting, cash-flow planning, and financial discipline.",
                "Budget packages, limits, plan-vs-actual; multi-scenario cash-flow forecasts "
                "and monthly payment planning.",
                "Configured financial reports in QAD; improved management visibility.",
                "Quarterly/annual IFRS/CAS reporting and external audit; compliance with legal "
                "documentation and HQ requirements.",
            ],
        },
        {
            "role": "Executive Director",
            "company": "Engineering Solutions (China / Russia)",
            "period": "2013 — 2017",
            "intro": "",
            "bullets": [
                "Managed procurement, sales, and foreign trade; owned sales planning and "
                "commercial execution.",
                "Led B2B contract negotiations, pricing, and profitability management in "
                "international trading.",
            ],
        },
        {
            "role": "Executive Director",
            "company": "New Technologies (Moscow, Russia)",
            "period": "2008 — 2013",
            "intro": "",
            "bullets": [
                "Ran sales, procurement, and operations with P&L accountability; pricing, "
                "purchasing, and key account relationships.",
                "Cross-border sourcing (Southeast Asia), licensing, and market entry.",
            ],
        },
        {
            "role": "Data Analyst / Analytics Engineer",
            "company": "Yofi Inc. (USA, remote)",
            "period": "February 2022 — October 2025",
            "intro": (
                "SaaS anti-fraud platform for Shopify merchants (enterprise customers "
                "including Lululemon)."
            ),
            "bullets": [
                "Supported ROI and operational dashboards; merchant KPIs, billing, and "
                "commercial performance on a high-growth tech platform.",
                "Partnered with product/commercial stakeholders on metric definitions and "
                "reporting for strategic decisions.",
                "Remote international SaaS environment with US-based teams.",
            ],
        },
        {
            "role": "Business Analyst / AI Platform",
            "company": "Sinoptics AI (remote)",
            "period": "March 2025 — Present",
            "intro": "",
            "bullets": [
                "Structured cross-functional requirements; validation rules improving "
                "reporting transparency.",
                "Coordinated pilots; executive-style status and risk summaries.",
            ],
        },
        {
            "role": "Head of IT and Finance / Head of IT Department",
            "company": "New Engineering Solution / Engineering Solutions LLC",
            "period": "2003 — 2013",
            "intro": "",
            "bullets": [
                "ERP and reporting systems, management accounting, financial control; "
                "international contracts, Incoterms, import/export supply chains.",
            ],
        },
    ]

    for exp in experiences:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{exp['role']} — ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(exp["company"])
        r2.bold = True
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)

        p = doc.add_paragraph()
        r = p.add_run(exp["period"])
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        if exp["intro"]:
            p = doc.add_paragraph(exp["intro"])
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(9.5)

        for bullet in exp["bullets"]:
            p = doc.add_paragraph(bullet, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(9.5)

    add_horizontal_line(doc)

    h = doc.add_heading("Education", level=1)
    h.runs[0].font.size = Pt(12)
    for item in [
        "MBA, Russian Presidential Academy (RANEPA), 2005–2007",
        "PhD in Physics and Mathematics, Southern Federal University, 2000–2003",
        "Physics Diploma (Specialist), Southern Federal University, 1994–1999",
        "College of Radio-Electronic Instrumentation, 1991–1995",
    ]:
        p = doc.add_paragraph(f"• {item}")
        for run in p.runs:
            run.font.size = Pt(10)

    add_horizontal_line(doc)

    h = doc.add_heading("Certifications", level=1)
    h.runs[0].font.size = Pt(12)
    for item in [
        "CAP — Certified Accountant Practitioner",
        "CPA Russia: IFRS reporting, management accounting, taxation, Russian legislation",
        "Chief Accountant Courses",
    ]:
        p = doc.add_paragraph(f"• {item}")
        for run in p.runs:
            run.font.size = Pt(10)

    add_horizontal_line(doc)

    h = doc.add_heading("Why Director, Deal Desk at DDN", level=1)
    h.runs[0].font.size = Pt(12)
    why = (
        "DDN operates where AI infrastructure, complex enterprise sales, and financial "
        "discipline must work together. My career combines CFO-level governance (approvals, "
        "controls, board reporting) with front-line commercial experience (B2B negotiation, "
        "pricing, profitability) and recent SaaS/tech platform exposure.\n\n"
        "I am motivated by building the function while delivering outcomes: formalizing "
        "playbooks, modeling exceptions with data, and acting as the air traffic controller "
        "between Sales, Finance, Legal, and Operations. Based in Russia (Remote), ready to "
        "own EMEA deal flow with independence and time-zone resilience."
    )
    p = doc.add_paragraph(why)
    for run in p.runs:
        run.font.size = Pt(10)

    h = doc.add_heading("DDN Cultural Fit", level=1)
    h.runs[0].font.size = Pt(12)
    _add_table(
        doc,
        [
            (
                "Self-Starter",
                "Built finance functions from zero; proactively implemented controls without "
                "a mature playbook",
            ),
            (
                "Success Orientation",
                "Delivered budget cycles, audit-ready reporting, commercial outcomes under P&L",
            ),
            (
                "Problem Solving",
                "Systematic plan-vs-actual, liquidity, and deal profitability analysis",
            ),
            (
                "Innovative",
                "Improved QAD reporting, automated analytics, formalized approval workflows",
            ),
        ],
        col0_width=4,
        col1_width=13.5,
    )

    p = doc.add_paragraph()
    r = p.add_run(
        "Resume prepared specifically for the Director, Deal Desk position at "
        "DataDirect Networks (DDN)."
    )
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def main():
    output_dir = os.path.dirname(os.path.abspath(__file__))
    path = os.path.join(output_dir, "Director_Deal_Desk_DDN_Kurnosenko_EN.docx")
    create_resume().save(path)
    print(f"Saved: {path}")


if __name__ == "__main__":
    main()
