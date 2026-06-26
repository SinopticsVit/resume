#!/usr/bin/env python3
"""Generate DOCX for Plata Risk DWH Analytics Engineer resume and cover letter."""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).parent
RESUME_MD = HERE / "Analytics_Engineer_Plata_Risk_DWH_Kurnosenko_EN.md"
RESUME_DOCX = HERE / "Analytics_Engineer_Plata_Risk_DWH_Kurnosenko_EN.docx"
COVER_MD = HERE / "Cover_Letter_Analytics_Engineer_Plata_Risk_DWH_Kurnosenko_EN.md"
COVER_DOCX = HERE / "Cover_Letter_Analytics_Engineer_Plata_Risk_DWH_Kurnosenko_EN.docx"

NAVY = RGBColor(0x1F, 0x49, 0x7D)
STEEL = RGBColor(0x2E, 0x50, 0x90)
FONT = "Calibri"


def set_run_font(run, *, bold=False, italic=False, size=10, color=None):
    run.bold = bold
    run.italic = italic
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def add_rich_text(paragraph, text, *, size=10, bold_all=False):
    if bold_all:
        run = paragraph.add_run(text)
        set_run_font(run, bold=True, size=size)
        return
    pattern = re.compile(r"\*\*(.+?)\*\*")
    pos = 0
    if not pattern.search(text):
        run = paragraph.add_run(text)
        set_run_font(run, size=size)
        return
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size)
        run = paragraph.add_run(match.group(1))
        set_run_font(run, bold=True, size=size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size)


def add_horizontal_line(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for name in ("top", "left", "right"):
        edge = OxmlElement(f"w:{name}")
        edge.set(qn("w:val"), "nil")
        borders.append(edge)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "2E5090")
    borders.append(bottom)
    tc_pr.append(borders)


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_markdown_table(doc, table_lines):
    rows = []
    for line in table_lines:
        if re.match(r"^\|[-| :]+\|$", line.strip()):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:
        return

    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    widths = [Cm(5.2), Cm(10.8)] if col_count == 2 else None
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            add_rich_text(p, text, size=8.5 if r_idx else 9, bold_all=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, "1F497D")
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if widths and c_idx < len(widths):
                cell.width = widths[c_idx]


def convert_resume_md_to_docx(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    i = 0
    header_done = False

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            add_horizontal_line(doc)
            if not header_done:
                header_done = True
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:])
            set_run_font(run, bold=True, size=18, color=NAVY)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if stripped.startswith("## "):
            h = doc.add_heading(stripped[3:], level=1)
            h.runs[0].font.size = Pt(12)
            h.runs[0].font.color.rgb = NAVY
            i += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph()
            add_rich_text(p, stripped[4:], size=10.5)
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = STEEL
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.5)
            add_rich_text(p, stripped[2:])
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and stripped.count("*") == 2:
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            set_run_font(run, italic=True, size=9.5, color=RGBColor(0x55, 0x55, 0x55))
            i += 1
            continue

        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:-2])
            set_run_font(run, bold=True, size=10.5, color=STEEL)
            i += 1
            continue

        p = doc.add_paragraph()
        if not header_done:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_rich_text(p, stripped, size=9 if not stripped.startswith("**") else 10)
        else:
            add_rich_text(p, stripped)
        i += 1

    doc.save(docx_path)


def convert_cover_md_to_docx(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)

    in_body = False
    skip_next_blank_after_hr = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# Cover Letter"):
            continue
        if stripped == "---":
            skip_next_blank_after_hr = True
            continue
        if stripped.startswith("**Vitaly Kurnosenko**"):
            p = doc.add_paragraph()
            run = p.add_run("Vitaly Kurnosenko")
            set_run_font(run, bold=True, size=16, color=NAVY)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if stripped.startswith("Shanghai") or stripped.startswith("+86"):
            p = doc.add_paragraph()
            add_rich_text(p, stripped.replace("**", ""), size=9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if stripped.startswith("**June") or stripped.startswith("**Hiring") or stripped.startswith("**DWH"):
            p = doc.add_paragraph()
            add_rich_text(p, stripped.replace("**", ""), size=11)
            if "Hiring" in stripped:
                p.runs[0].bold = True
            continue
        if stripped == "Banco Plata":
            p = doc.add_paragraph("Banco Plata")
            continue
        if stripped.startswith("Dear "):
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_after = Pt(8)
            in_body = True
            continue
        if stripped.startswith("Sincerely"):
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_before = Pt(12)
            in_body = False
            continue
        if stripped == "**Vitaly Kurnosenko**":
            p = doc.add_paragraph()
            run = p.add_run("Vitaly Kurnosenko")
            set_run_font(run, bold=True, size=11)
            continue
        if in_body:
            p = doc.add_paragraph()
            add_rich_text(p, stripped, size=11)
            p.paragraph_format.space_after = Pt(8)

    doc.save(docx_path)


if __name__ == "__main__":
    convert_resume_md_to_docx(RESUME_MD, RESUME_DOCX)
    print(f"DOCX written: {RESUME_DOCX}")
    convert_cover_md_to_docx(COVER_MD, COVER_DOCX)
    print(f"DOCX written: {COVER_DOCX}")
