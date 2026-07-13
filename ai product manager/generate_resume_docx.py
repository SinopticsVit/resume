#!/usr/bin/env python3
"""Generate DOCX for AI Product Manager resume and cover letter."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

HERE = Path(__file__).parent
sys.path.insert(0, str(HERE.parent / "analytics_engineer"))

from generate_plata_risk_docx import (
    NAVY,
    add_rich_text,
    convert_resume_md_to_docx,
    set_run_font,
)

MD = HERE / "AI_Product_Manager_Kurnosenko_EN.md"
DOCX = HERE / "AI_Product_Manager_Kurnosenko_EN.docx"
COVER_MD = HERE / "Cover_Letter_AI_Product_Manager_Kurnosenko_EN.md"
COVER_DOCX = HERE / "Cover_Letter_AI_Product_Manager_Kurnosenko_EN.docx"


def convert_cover_md_to_docx(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    in_body = False

    for line in lines:
        stripped = line.strip()
        if not stripped or stripped.startswith("# ") or stripped == "---":
            continue

        if stripped == "**Vitaly Kurnosenko**" and not in_body:
            p = doc.add_paragraph()
            run = p.add_run("Vitaly Kurnosenko")
            set_run_font(run, bold=True, size=16, color=NAVY)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if stripped.startswith("Shanghai") or stripped.startswith("+86"):
            p = doc.add_paragraph()
            add_rich_text(p, stripped.replace("**", ""), size=9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if re.match(r"^\*\*.+\*\*$", stripped) and not in_body and not stripped.startswith("Dear"):
            p = doc.add_paragraph()
            text = stripped.strip("*")
            run = p.add_run(text)
            set_run_font(run, bold=("Hiring" in text), size=11)
            continue

        if not in_body and not stripped.startswith("Dear") and not stripped.startswith("Sincerely"):
            p = doc.add_paragraph(stripped)
            continue

        if stripped.startswith("Dear "):
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_after = Pt(8)
            in_body = True
            continue

        if stripped.startswith("Sincerely"):
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_before = Pt(12)
            in_body = False
            continue

        if stripped == "**Vitaly Kurnosenko**":
            p = doc.add_paragraph()
            run = p.add_run("Vitaly Kurnosenko")
            set_run_font(run, bold=True, size=11)
            continue

        if in_body:
            p = doc.add_paragraph()
            add_rich_text(p, stripped, size=11)
            p.paragraph_format.space_after = Pt(8)

    doc.save(docx_path)


if __name__ == "__main__":
    convert_resume_md_to_docx(MD, DOCX)
    print("DOCX written:", DOCX)
    convert_cover_md_to_docx(COVER_MD, COVER_DOCX)
    print("DOCX written:", COVER_DOCX)
