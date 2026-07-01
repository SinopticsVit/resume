#!/usr/bin/env python3
"""Generate DOCX from Antifraud_Analyst_Senior_Kurnosenko_{EN,RU}.md."""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).parent
RESUME_VARIANTS = (
    ("Antifraud_Analyst_Senior_Kurnosenko_EN.md", "Antifraud_Analyst_Senior_Kurnosenko_EN.docx"),
    ("Antifraud_Analyst_Senior_Kurnosenko_RU.md", "Antifraud_Analyst_Senior_Kurnosenko_RU.docx"),
)

NAVY = RGBColor(0x1F, 0x49, 0x7D)
STEEL = RGBColor(0x2E, 0x50, 0x90)
FONT = "Calibri"


def set_run_font(run, *, bold=False, italic=False, size=10, color=None):
    run.bold = bold
    run.italic = italic
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def add_horizontal_line(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for name in ("top", "left", "right"):
        edge = OxmlElement(f"w:{name}")
        edge.set(qn("w:val"), "nil")
        borders.append(edge)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "2E5090")
    borders.append(bottom)
    tc_pr.append(borders)
    tr = table.rows[0]._tr
    tr_pr = tr.get_or_add_trPr()
    height = OxmlElement("w:trHeight")
    height.set(qn("w:val"), "50")
    height.set(qn("w:hRule"), "exact")
    tr_pr.append(height)


def add_rich_paragraph(doc, text, *, style=None, bullet=False, center=False, size=10):
    p = doc.add_paragraph(style="List Bullet" if bullet else style)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if bullet:
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.25)

    pattern = re.compile(r"\*\*(.+?)\*\*")
    pos = 0
    if not pattern.search(text):
        run = p.add_run(text)
        set_run_font(run, size=size)
        return p

    for match in pattern.finditer(text):
        if match.start() > pos:
            run = p.add_run(text[pos : match.start()])
            set_run_font(run, size=size)
        run = p.add_run(match.group(1))
        set_run_font(run, bold=True, size=size)
        pos = match.end()
    if pos < len(text):
        run = p.add_run(text[pos:])
        set_run_font(run, size=size)
    return p


def add_section_heading(doc, title):
    h = doc.add_heading(title, level=1)
    h.runs[0].font.size = Pt(12)
    h.runs[0].font.color.rgb = NAVY
    h.runs[0].font.name = FONT
    h.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def add_job_heading(doc, title):
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run_font(run, bold=True, size=10.5, color=STEEL)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)


def convert_md_to_docx(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    i = 0
    header_done = False

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            if not header_done:
                add_horizontal_line(doc)
                header_done = True
            else:
                add_horizontal_line(doc)
            i += 1
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:])
            set_run_font(run, bold=True, size=18, color=NAVY)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if stripped.startswith("## "):
            add_section_heading(doc, stripped[3:])
            i += 1
            continue

        if stripped.startswith("### "):
            add_job_heading(doc, stripped[4:])
            i += 1
            continue

        if stripped.startswith("- "):
            add_rich_paragraph(doc, stripped[2:], bullet=True)
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and stripped.count("*") == 2:
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            set_run_font(run, italic=True, size=9.5, color=RGBColor(0x55, 0x55, 0x55))
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:-2])
            set_run_font(run, bold=True, size=10.5, color=STEEL)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        if not header_done and stripped.startswith("**"):
            add_rich_paragraph(doc, stripped, center=True, size=10 if "|" in stripped else 9)
            i += 1
            continue

        if not header_done:
            add_rich_paragraph(doc, stripped, center=True, size=9)
            i += 1
            continue

        add_rich_paragraph(doc, stripped)
        i += 1

    doc.save(docx_path)


if __name__ == "__main__":
    for md_name, docx_name in RESUME_VARIANTS:
        md_path = HERE / md_name
        docx_path = HERE / docx_name
        convert_md_to_docx(md_path, docx_path)
        print(f"DOCX written: {docx_path}")
