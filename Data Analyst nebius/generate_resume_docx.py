#!/usr/bin/env python3
"""Generate professional DOCX resume for Data Analyst at Nebius (EN / RU)."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_horizontal_line(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for border_name in ["top", "left", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "nil")
        tcBorders.append(border)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "2E5090")
    tcBorders.append(bottom)
    tcPr.append(tcBorders)
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), "50")
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


def _content(lang):
    if lang == "RU":
        return {
            "name": "Курносенко Виталий Николаевич",
            "title": (
                "Аналитик данных | SQL · Python · dbt · Airflow · BigQuery | "
                "Дашборды · Качество данных · ETL/ELT"
            ),
            "contact": (
                "Локация: Азия (готов к релокации)  |  Гражданство: Россия  |  "
                "Телефон: +86 15601694273  |  Email: kurnosenko@mail.ru  |  "
                "Telegram: @vitaly_kur  |  WeChat: porohnya"
            ),
            "languages": (
                "Языки: русский — родной; английский — B2 рабочий уровень; "
                "китайский (мандарин) — HSK 4"
            ),
            "summary_title": "Профессиональное резюме",
            "summary": (
                "Аналитик данных с 3+ годами практического опыта на высоконагруженной "
                "e-commerce anti-fraud платформе (Yofi, США — мерчанты Shopify, включая "
                "enterprise retail). Перевожу бизнес- и стейкхолдерские вопросы в надёжные "
                "датасеты, дашборды и прикладные инсайты с помощью SQL, Python, dbt, "
                "Airflow и BigQuery.\n\n"
                "Сильная база в математической статистике и теории вероятностей "
                "(канд. физ.-мат. наук) — применяю в описательном анализе, выявлении "
                "аномалий, логике воронок и когорт, аккуратной интерпретации метрик. "
                "Комфортно работаю end-to-end: валидация источников, превращение сырых "
                "данных в доверенные витрины, визуализации для стейкхолдеров, понятные "
                "root-cause summary при сбоях данных или пайплайнов.\n\n"
                "Опыт работы в международных удалённых командах; готов поддерживать "
                "ежедневные решения, улучшать проверки качества данных и участвовать в "
                "AI-assisted analytics в Nebius."
            ),
            "fit_title": "Соответствие вакансии — Data Analyst в Nebius",
            "fit_rows": [
                (
                    "Data Analyst или аналогичная роль",
                    "3+ года как Data Analyst / Analytics Engineer на продакшн data platform; "
                    "ранее — продуктовая и финансовая аналитика",
                ),
                (
                    "Сильный Python и SQL",
                    "Продвинутый SQL (CTE, оконные функции) на BigQuery, PostgreSQL, MS SQL; "
                    "Python (pandas, PySpark) для анализа и поддержки пайплайнов",
                ),
                (
                    "Статистика и аналитическое мышление",
                    "Канд. физ.-мат. наук; описательная статистика, интерпретация с учётом "
                    "гипотез; анализ паттернов fraud и конверсии",
                ),
                (
                    "Дашборды и коммуникация",
                    "ROI- и операционные дашборды; отчётность для стейкхолдеров; "
                    "документированные метрики и бизнес-правила",
                ),
                (
                    "Качество данных и детали",
                    "dbt tests, валидация пайплайнов, исправление null/несогласованностей, "
                    "согласованные определения метрик",
                ),
                (
                    "Бизнес → структурированный анализ",
                    "Сбор требований с product и аналитиками; ad-hoc → переиспользуемые "
                    "витрины и отчёты",
                ),
                (
                    "BI-инструменты",
                    "Power BI; готов работать в Looker / Superset / Tableau по стандартам команды",
                ),
                (
                    "Modern data stack",
                    "dbt на BigQuery; ~25 Airflow DAG; PostgreSQL, MongoDB, Spanner как источники",
                ),
                (
                    "Когорты / воронки / эксперименты",
                    "Сессионизация, логика воронок и merchant KPI; оценка качества AI-workflow",
                ),
                (
                    "Облако / технический продукт",
                    "AWS и GCP data platforms; high-volume events; ML-adjacent feature analytics",
                ),
            ],
            "skills_title": "Ключевые навыки",
            "skills": [
                (
                    "Анализ данных и SQL",
                    "Продвинутый SQL — CTE, оконные функции, агрегации; BigQuery, PostgreSQL, "
                    "MS SQL, Spanner; оптимизация запросов; exploratory analysis",
                ),
                (
                    "Python и автоматизация",
                    "pandas, pandas-gbq, PySpark; исследование данных, backfill, "
                    "автоматизация отчётности; Excel когда эффективнее",
                ),
                (
                    "Data platform и пайплайны",
                    "dbt, Airflow, Airbyte, Spark на Kubernetes; lakehouse на GCS / BigLake",
                ),
                (
                    "Качество данных и документация",
                    "Автотесты, валидация источников, мониторинг freshness; "
                    "внутренняя документация; code review",
                ),
                (
                    "Визуализация и отчётность",
                    "Power BI; операционные и ROI-дашборды; понятные рекомендации",
                ),
                (
                    "Домен и методы",
                    "E-commerce anti-fraud — заказы, billing, customer clusters, merchant KPI; "
                    "воронки, когорты, anomaly thinking; финансовая дисциплина метрик",
                ),
                (
                    "Коллаборация",
                    "Product, engineering, аналитики; согласование KPI; "
                    "международные удалённые команды",
                ),
            ],
            "exp_title": "Опыт работы",
            "experiences": [
                {
                    "role": "Data Analyst / Analytics Engineer",
                    "company": "Yofi Inc. (США, удалённо)",
                    "period": "февраль 2022 — октябрь 2025",
                    "intro": (
                        "Anti-fraud и customer-intelligence платформа для мерчантов Shopify "
                        "(enterprise-клиенты, включая Lululemon)."
                    ),
                    "bullets": [
                        "Создавал и поддерживал analytical datasets и витрины в BigQuery с dbt — "
                        "измерения заказов и клиентов, fraud indicators, merchant reporting layers; "
                        "стандарты качества SQL и документированная бизнес-логика.",
                        "Поддержка ETL/ELT: оркестрация ~25 продакшн Airflow DAG (Spark jobs, "
                        "Airbyte syncs, incremental и full-refresh); мониторинг сбоев и здоровья пайплайнов.",
                        "Оптимизированный SQL для сложного анализа — join MongoDB, PostgreSQL, "
                        "Spanner и DWH; оконные функции и CTE для сессионизации, воронок, "
                        "risk-confirmation logic.",
                        "Качество данных через автотесты, batch-валидацию, исправление null и "
                        "schema drift; согласованные relational schemas.",
                        "PySpark-нагрузки для lakehouse ingestion на GCS; настройка partitioning "
                        "для high-volume event processing.",
                        "ROI- и операционные дашборды; Python/pandas для ad-hoc и автоматизации отчётов.",
                        "Документирование моделей и метрик; обучение аналитиков dbt, Airbyte, Spark.",
                        "Feature analytics и realtime severity workflows для risk reporting.",
                        "Интеграция Shopify, partner webhooks и маркетинговых источников через Airbyte.",
                    ],
                    "stack": (
                        "Стек: SQL, Python, PySpark, dbt, Airflow, Airbyte, BigQuery, GCS, "
                        "PostgreSQL, MongoDB, Redis, Neo4j, Spanner, Power BI, AWS, GCP."
                    ),
                },
                {
                    "role": "Business Analyst / AI Platform",
                    "company": "Sinoptics AI (удалённо)",
                    "period": "март 2025 — настоящее время",
                    "intro": "",
                    "bullets": [
                        "Структурирование требований для document-processing и analytics workflows; "
                        "validation rules улучшили прозрачность отчётности.",
                        "Data flows на Python и SQL для product analytics и AI-assisted features.",
                        "Оценка AI workflows — quality checks и итерации по feedback.",
                        "Координация пилотов; понятная status-отчётность для нетехнической аудитории.",
                    ],
                    "stack": "",
                },
                {
                    "role": "Руководитель IT и финансов",
                    "company": "Engineering Solutions LLC",
                    "period": "март 2013 — декабрь 2017",
                    "intro": "",
                    "bullets": [
                        "Финансовые и IT-операции: инвестиционный анализ, бюджетирование, "
                        "управленческая отчётность и автоматизация для KPI tracking.",
                    ],
                    "stack": "",
                },
                {
                    "role": "Руководитель IT-отдела",
                    "company": "New Engineering Solution",
                    "period": "апрель 2003 — февраль 2013",
                    "intro": "",
                    "bullets": [
                        "ERP и reporting systems; структурированная внутренняя отчётность "
                        "для управленческих решений.",
                    ],
                    "stack": "",
                },
            ],
            "edu_title": "Образование",
            "education": [
                "Кандидат физико-математических наук, Южный федеральный университет, 2000–2003",
                "Диплом специалиста по физике, Южный федеральный университет, 1994–1999",
                "MBA, РАНХиГС, 2005–2007",
                "Колледж радиоэлектронного приборостроения, 1991–1995",
            ],
            "cert_title": "Сертификаты",
            "certifications": [
                "CAP — Certified Accountant Practitioner",
                "Курсы CPA Russia: финансовый учёт и МСФО, управленческий учёт, налогообложение",
                "Курсы главного бухгалтера",
            ],
            "why_title": "Почему Data Analyst в Nebius",
            "why": (
                "Nebius находится на стыке облачной инфраструктуры, AI и data-driven "
                "decision-making — контекст, где аналитическая строгость и умение объяснять "
                "trade-offs не менее важны, чем инструменты. Мой бэкграунд сочетает "
                "продакшн-опыт на modern data stack (dbt, Airflow, BigQuery, Python, SQL) "
                "с PhD-уровнем в статистике и годами KPI-driven reporting в регулируемых средах.\n\n"
                "Комфортно владею полным циклом: уточнить бизнес-вопрос, провалидировать "
                "определения, сформировать доверенные датасеты, построить дашборды, которыми "
                "реально пользуются, и копнуть глубже, когда цифры выглядят подозрительно. "
                "Готов поддерживать качество ETL, anomaly-oriented analysis и AI-assisted "
                "analytics по мере масштабирования Nebius."
            ),
            "footer": "Резюме подготовлено специально для позиции Data Analyst в Nebius.",
        }

    return {
        "name": "Kurnosenko Vitaly Nikolaevich",
        "title": (
            "Data Analyst | SQL · Python · dbt · Airflow · BigQuery | "
            "Dashboards · Data Quality · ETL/ELT"
        ),
        "contact": (
            "Location: Asia (open to relocation)  |  Citizenship: Russia  |  "
            "Phone: +86 15601694273  |  Email: kurnosenko@mail.ru  |  "
            "Telegram: @vitaly_kur  |  WeChat: porohnya"
        ),
        "languages": (
            "Languages: Russian — native; English — B2 working proficiency; "
            "Chinese (Mandarin) — HSK 4"
        ),
        "summary_title": "Professional Summary",
        "summary": (
            "Data analyst with 3+ years of hands-on experience on a high-volume e-commerce "
            "anti-fraud platform (Yofi, USA — Shopify merchants including enterprise retail). "
            "Translate business and stakeholder questions into reliable datasets, dashboards, "
            "and actionable insights using SQL, Python, dbt, Airflow, and BigQuery.\n\n"
            "Strong grounding in mathematical statistics and probability (PhD in Physics and "
            "Mathematics) — applied to descriptive analysis, anomaly detection, funnel and "
            "cohort logic, and careful interpretation of metrics. Comfortable end-to-end: "
            "validating sources, shaping raw data into trusted marts, building visualizations "
            "stakeholders rely on, and writing clear root-cause summaries when data or "
            "pipelines break.\n\n"
            "Experienced in international remote teams; ready to support day-to-day "
            "decision-making, improve data quality checks, and contribute to AI-assisted "
            "analytics workflows at Nebius."
        ),
        "fit_title": "Role Fit — Data Analyst at Nebius",
        "fit_rows": [
            (
                "Data Analyst or similar role",
                "3+ years as Data Analyst / Analytics Engineer on production data platform; "
                "earlier product and finance analytics background",
            ),
            (
                "Strong Python and SQL",
                "Advanced SQL (CTEs, window functions) on BigQuery, PostgreSQL, MS SQL; "
                "Python (pandas, PySpark) for analysis and pipeline support",
            ),
            (
                "Statistics and analytical thinking",
                "PhD in mathematical sciences; descriptive statistics, hypothesis-aware "
                "interpretation; fraud and conversion pattern analysis",
            ),
            (
                "Dashboards and communication",
                "ROI and operational dashboards; stakeholder reporting; documented metrics "
                "and business rules",
            ),
            (
                "Data quality and detail",
                "dbt tests, pipeline validation, null/consistency fixes, aligned metric "
                "definitions across teams",
            ),
            (
                "Business → structured analysis",
                "Requirements with product and analysts; recurring ad-hoc work into reusable "
                "marts and reports",
            ),
            (
                "BI tools",
                "Power BI; ready for Looker / Superset / Tableau per team standards",
            ),
            (
                "Modern data stack",
                "dbt on BigQuery; ~25 Airflow DAGs; PostgreSQL, MongoDB, Spanner as sources",
            ),
            (
                "Cohort / funnel / experiments",
                "Sessionization, funnel and merchant KPI logic; AI workflow quality evaluation",
            ),
            (
                "Cloud / technical product",
                "AWS and GCP data platforms; high-volume events; ML-adjacent feature analytics",
            ),
        ],
        "skills_title": "Core Skills",
        "skills": [
            (
                "Data Analysis & SQL",
                "Advanced SQL — CTEs, window functions, aggregations; BigQuery, PostgreSQL, "
                "MS SQL, Spanner; query tuning; exploratory analysis",
            ),
            (
                "Python & Automation",
                "pandas, pandas-gbq, PySpark; exploration, backfills, reporting automation; "
                "Excel when most efficient",
            ),
            (
                "Data Platform & Pipelines",
                "dbt, Airflow, Airbyte, Spark on Kubernetes; lakehouse on GCS / BigLake",
            ),
            (
                "Data Quality & Documentation",
                "Automated tests, source validation, freshness monitoring; internal "
                "documentation; code reviews",
            ),
            (
                "Visualization & Reporting",
                "Power BI; operational and ROI dashboards; plain-language recommendations",
            ),
            (
                "Domain & Methods",
                "E-commerce anti-fraud — orders, billing, customer clusters, merchant KPIs; "
                "funnel, cohort, anomaly thinking; finance metrics discipline",
            ),
            (
                "Collaboration",
                "Product, engineering, analysts; KPI definition alignment; international remote teams",
            ),
        ],
        "exp_title": "Work Experience",
        "experiences": [
            {
                "role": "Data Analyst / Analytics Engineer",
                "company": "Yofi Inc. (USA, remote)",
                "period": "February 2022 — October 2025",
                "intro": (
                    "Anti-fraud and customer-intelligence platform for Shopify merchants "
                    "(enterprise customers including Lululemon)."
                ),
                "bullets": [
                    "Built and maintained analytical datasets and marts in BigQuery using dbt — "
                    "order and customer dimensions, fraud indicators, merchant reporting layers; "
                    "enforced SQL quality standards and documented business logic.",
                    "Owned ETL/ELT support: orchestrated ~25 production Airflow DAGs (Spark jobs, "
                    "Airbyte syncs, incremental and full-refresh); monitored failures and pipeline health.",
                    "Wrote optimized SQL for complex analysis — joins across MongoDB, PostgreSQL, "
                    "Spanner and the warehouse; window functions and CTEs for sessionization, "
                    "funnels, and risk-confirmation logic.",
                    "Ensured data quality via automated tests, batch validation, and fixes for "
                    "nulls and schema drift; aligned relational schemas across application and analytics.",
                    "Developed PySpark workloads for lakehouse ingestion on GCS; tuned partitioning "
                    "for high-volume event processing.",
                    "Supported ROI and operational dashboards; Python/pandas for ad-hoc exploration "
                    "and report automation.",
                    "Documented models and metrics; trained analysts on dbt, Airbyte, and Spark practices.",
                    "Collaborated on feature analytics and realtime severity workflows for risk reporting.",
                    "Integrated Shopify, partner webhooks, and marketing sources into the analytical lake via Airbyte.",
                ],
                "stack": (
                    "Stack: SQL, Python, PySpark, dbt, Airflow, Airbyte, BigQuery, GCS, "
                    "PostgreSQL, MongoDB, Redis, Neo4j, Spanner, Power BI, AWS, GCP."
                ),
            },
            {
                "role": "Business Analyst / AI Platform",
                "company": "Sinoptics AI (remote)",
                "period": "March 2025 — Present",
                "intro": "",
                "bullets": [
                    "Structured requirements for document-processing and analytics workflows; "
                    "validation rules improved reporting transparency.",
                    "Built data flows in Python and SQL for product analytics and AI-assisted features.",
                    "Participated in evaluation of AI workflows — quality checks and iteration from feedback.",
                    "Coordinated pilots; prepared clear status reporting for non-technical stakeholders.",
                ],
                "stack": "",
            },
            {
                "role": "Head of IT and Finance",
                "company": "Engineering Solutions LLC",
                "period": "March 2013 — December 2017",
                "intro": "",
                "bullets": [
                    "Led financial and IT operations: investment analysis, budgeting, management "
                    "reporting, and automation supporting KPI tracking.",
                ],
                "stack": "",
            },
            {
                "role": "Head of IT Department",
                "company": "New Engineering Solution",
                "period": "April 2003 — February 2013",
                "intro": "",
                "bullets": [
                    "Owned ERP and reporting systems; structured internal reporting for management decisions.",
                ],
                "stack": "",
            },
        ],
        "edu_title": "Education",
        "education": [
            "PhD in Physics and Mathematics, Southern Federal University, 2000–2003",
            "Physics Diploma (Specialist), Southern Federal University, 1994–1999",
            "MBA, RANEPA, 2005–2007",
            "College of Radio-Electronic Instrumentation, 1991–1995",
        ],
        "cert_title": "Certifications",
        "certifications": [
            "CAP — Certified Accountant Practitioner",
            "CPA Russia: financial accounting and IFRS reporting, management accounting, taxation",
            "Chief Accountant Courses",
        ],
        "why_title": "Why Data Analyst at Nebius",
        "why": (
            "Nebius sits at the intersection of cloud infrastructure, AI, and data-driven "
            "decision-making — where analytical rigor and the ability to explain trade-offs "
            "matter as much as tooling. My background combines production experience on a modern "
            "data stack (dbt, Airflow, BigQuery, Python, SQL) with a PhD-level foundation in "
            "statistics and years of KPI-driven reporting in regulated environments.\n\n"
            "I am comfortable owning the full loop: clarify the business question, validate "
            "definitions, shape trustworthy datasets, build dashboards people use, and dig deeper "
            "when numbers look wrong. Ready to support ETL quality, anomaly-oriented analysis, "
            "and AI-assisted analytics as Nebius scales."
        ),
        "footer": "Resume prepared specifically for the Data Analyst position at Nebius.",
    }


def create_resume(lang="EN"):
    c = _content(lang)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    p = doc.add_paragraph()
    run = p.add_run(c["name"])
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(c["title"])
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    run = p.add_run(c["contact"])
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(c["languages"])
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    add_horizontal_line(doc)

    h = doc.add_heading(c["summary_title"], level=1)
    h.runs[0].font.size = Pt(12)
    p = doc.add_paragraph(c["summary"])
    for run in p.runs:
        run.font.size = Pt(10)
    add_horizontal_line(doc)

    h = doc.add_heading(c["fit_title"], level=1)
    h.runs[0].font.size = Pt(12)
    table = doc.add_table(rows=len(c["fit_rows"]), cols=2)
    table.style = "Table Grid"
    for i, (req, evidence) in enumerate(c["fit_rows"]):
        c0 = table.rows[i].cells[0]
        c0.text = req
        for para in c0.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(c0, "E8EEF4")
        c1 = table.rows[i].cells[1]
        c1.text = evidence
        for para in c1.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
    for row in table.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(12)
    doc.add_paragraph()
    add_horizontal_line(doc)

    h = doc.add_heading(c["skills_title"], level=1)
    h.runs[0].font.size = Pt(12)
    table = doc.add_table(rows=len(c["skills"]), cols=2)
    table.style = "Table Grid"
    for i, (area, detail) in enumerate(c["skills"]):
        c0 = table.rows[i].cells[0]
        c0.text = area
        for para in c0.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(c0, "E8EEF4")
        c1 = table.rows[i].cells[1]
        c1.text = detail
        for para in c1.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
    doc.add_paragraph()
    add_horizontal_line(doc)

    h = doc.add_heading(c["exp_title"], level=1)
    h.runs[0].font.size = Pt(12)
    for exp in c["experiences"]:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{exp['role']} — ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(exp["company"])
        r2.bold = True
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)

        p = doc.add_paragraph()
        r = p.add_run(exp["period"])
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        if exp["intro"]:
            p = doc.add_paragraph(exp["intro"])
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(9.5)

        for bullet in exp["bullets"]:
            p = doc.add_paragraph(bullet, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(9.5)

        if exp["stack"]:
            p = doc.add_paragraph(exp["stack"])
            for run in p.runs:
                run.font.size = Pt(9)
                run.italic = True

    add_horizontal_line(doc)

    h = doc.add_heading(c["edu_title"], level=1)
    h.runs[0].font.size = Pt(12)
    for item in c["education"]:
        p = doc.add_paragraph(f"• {item}")
        for run in p.runs:
            run.font.size = Pt(10)

    add_horizontal_line(doc)

    h = doc.add_heading(c["cert_title"], level=1)
    h.runs[0].font.size = Pt(12)
    for item in c["certifications"]:
        p = doc.add_paragraph(f"• {item}")
        for run in p.runs:
            run.font.size = Pt(10)

    add_horizontal_line(doc)

    h = doc.add_heading(c["why_title"], level=1)
    h.runs[0].font.size = Pt(12)
    p = doc.add_paragraph(c["why"])
    for run in p.runs:
        run.font.size = Pt(10)

    p = doc.add_paragraph()
    r = p.add_run(c["footer"])
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def main():
    output_dir = os.path.dirname(os.path.abspath(__file__))

    for lang, filename in (
        ("EN", "Data_Analyst_Kurnosenko_EN.docx"),
        ("RU", "Data_Analyst_Kurnosenko_RU.docx"),
    ):
        path = os.path.join(output_dir, filename)
        create_resume(lang=lang).save(path)
        print(f"Saved: {path}")


if __name__ == "__main__":
    main()
