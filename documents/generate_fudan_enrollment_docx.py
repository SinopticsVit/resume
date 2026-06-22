#!/usr/bin/env python3
"""Generate Fudan University enrollment / visa extension letter as DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).resolve().parent
OUT = HERE / "Fudan_Enrollment_Letter_Kurnosenko_Darya_EN.docx"

NAVY = RGBColor(0x1A, 0x3A, 0x6B)
MUTED = RGBColor(0x5A, 0x5A, 0x7A)


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_centered(doc: Document, text: str, *, bold: bool = False, size: int = 11,
                 color: RGBColor | None = None, space_after: int = 2) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)


def add_body(doc: Document, text: str, *, bold: bool = False, space_after: int = 6,
             justify: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)


def add_mixed_paragraph(doc: Document, parts: list[tuple[str, bool]], *,
                        space_after: int = 6, justify: bool = False) -> None:
    p = doc.add_paragraph()
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_info_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        cell_l = table.rows[i].cells[0]
        cell_r = table.rows[i].cells[1]
        cell_l.text = label
        cell_r.text = value
        for cell in (cell_l, cell_r):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
        for run in cell_l.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = MUTED
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)


def build() -> None:
    doc = Document()
    set_margins(doc)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    add_centered(doc, "FUDAN UNIVERSITY", bold=True, size=16, color=NAVY, space_after=2)
    for line in [
        "International Cultural Exchange School (ICES)",
        "220 Handan Road, Yangpu District, Shanghai 200433, P.R. China",
        "Tel: +86 21 6564 2292  |  Email: ices@fudan.edu.cn",
        "Website: https://www.ices.fudan.edu.cn",
    ]:
        add_centered(doc, line, size=9, color=MUTED, space_after=1)

    add_centered(
        doc,
        "CERTIFICATE OF ENROLLMENT AND VISA EXTENSION REQUEST",
        bold=True,
        size=13,
        color=NAVY,
        space_after=10,
    )

    add_mixed_paragraph(doc, [
        ("Document No.: ", True), ("FDU-ICES-2026-CL-0847", False),
        ("     Date of Issue: ", True), ("June 22, 2026", False),
    ], space_after=10)

    add_mixed_paragraph(doc, [("To:", True)], space_after=2)
    add_body(doc, "Shanghai Exit-Entry Administration", space_after=0)
    add_body(doc, "General Administration of Immigration, P.R. China", space_after=0)
    add_body(doc, "(or the relevant visa-issuing authority)", space_after=10)

    add_mixed_paragraph(doc, [
        ("Subject: ", True),
        ("Certificate of Enrollment — Request for Visa Extension", False),
    ], space_after=2)
    add_mixed_paragraph(doc, [
        ("Re: ", True),
        ("Ms. ", False), ("Darya Vitalievna Kurnosenko", True),
        (" (Russian Federation)", False),
    ], space_after=10)

    add_body(doc, "Dear Sir or Madam,", space_after=8)
    add_mixed_paragraph(doc, [
        ("This letter is issued by ", False),
        ("Fudan University", True),
        (", Shanghai, to confirm the enrollment status of the above-named student and "
         "to respectfully request an extension of her stay (visa/residence permit validity) "
         "in the People's Republic of China.", False),
    ], justify=True, space_after=10)

    add_heading(doc, "1. Student Information")
    add_info_table(doc, [
        ("Full Name", "KURNOSENKO, Darya Vitalievna"),
        ("Nationality", "Russian Federation"),
        ("Date of Birth", "May 23, 2008"),
        ("Passport No.", "[to be filled in]"),
    ])

    add_heading(doc, "2. Program of Study")
    add_mixed_paragraph(doc, [
        ("The student is currently enrolled in the ", False),
        ("Intensive Chinese Language Program", True),
        (" at the ", False),
        ("International Cultural Exchange School (ICES), Fudan University", True),
        (", Shanghai.", False),
    ], justify=True, space_after=6)
    add_info_table(doc, [
        ("Institution", "Fudan University"),
        ("Program", "Intensive Chinese Language Course"),
        ("Campus / Teaching Location", "Fudan University, Yangpu Campus, Shanghai"),
        ("Student Status", "Active — currently attending classes"),
        ("Last Day of Instruction", "July 3, 2026 (Friday)"),
    ])
    add_mixed_paragraph(doc, [
        ("Upon completion of the final class on ", False),
        ("July 3, 2026", True),
        (", the student will have fulfilled all academic requirements of the current "
         "language course. No further classes are scheduled after that date.", False),
    ], justify=True, space_after=10)

    add_heading(doc, "3. Request for Visa Extension")
    add_mixed_paragraph(doc, [
        ("We kindly request that the relevant authority ", False),
        ("extend the validity of the student's visa / residence permit until "
         "July 5, 2026 (inclusive)", True),
        (".", False),
    ], justify=True, space_after=6)
    add_body(
        doc,
        "This extension is necessary to allow the student sufficient time to:",
        space_after=4,
    )
    add_bullet(doc,
               "complete course-related formalities and personal departure preparations "
               "after the last day of study;")
    add_bullet(doc, "check out of accommodation and travel to the airport;")
    add_bullet(doc,
               "depart the People's Republic of China and return to the Russian Federation "
               "on July 5, 2026.")
    add_mixed_paragraph(doc, [
        ("The student's planned date of departure from China is ", False),
        ("July 5, 2026", True),
        (". She does not intend to remain in China beyond that date.", False),
    ], justify=True, space_after=10)

    add_heading(doc, "4. Confirmation")
    add_mixed_paragraph(doc, [
        ("Fudan University confirms that ", False),
        ("Ms. Darya Vitalievna Kurnosenko", True),
        (" is a bona fide student in good standing at our institution and is attending "
         "the Chinese language program described above.", False),
    ], justify=True, space_after=6)
    add_body(
        doc,
        "Should you require any further information, please do not hesitate to contact "
        "the International Cultural Exchange School at the address or telephone number "
        "indicated above.",
        justify=True,
        space_after=6,
    )
    add_body(doc, "Respectfully submitted for your consideration.", space_after=16)

    add_mixed_paragraph(doc, [
        ("For and on behalf of Fudan University", True),
    ], space_after=2)
    add_body(doc, "International Cultural Exchange School (ICES)", space_after=20)

    add_body(doc, "_______________________________", space_after=2)
    add_mixed_paragraph(doc, [
        ("Dr. [Name]", True),
    ], space_after=2)
    add_body(doc, "Director, International Cultural Exchange School", space_after=0)
    add_body(doc, "Fudan University", space_after=16)

    add_body(doc, "_______________________________", space_after=2)
    add_mixed_paragraph(doc, [("Official Seal", True)], space_after=2)
    add_body(doc, "Fudan University", space_after=16)

    p = doc.add_paragraph()
    run = p.add_run(
        "Note: This document is a draft mock-up (template) prepared for visa application "
        "purposes. It is not an official document issued by Fudan University. Placeholder "
        "fields (passport number, signatory name) must be completed before use. For an "
        "official certificate, contact Fudan University ICES directly."
    )
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT)
    print(f"[ok] {OUT}")


if __name__ == "__main__":
    build()
